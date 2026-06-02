#!/usr/bin/env python3
"""
Generate DevOps Documentation DOCX.
Reads all markdown files, applies the Tahoma-based styling from the
reference document, and writes DevOps_Documentation.docx.
"""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS = os.path.dirname(os.path.abspath(__file__))

FILES = [
    "README.md",
    "terraform.md",
    "docker.md",
    "kubernetes.md",
    "jenkins.md",
    "troubleshooting.md",
]

# ── colours (from reference doc) ─────────────────────────────────────────────
C_H1   = RGBColor(0x1A, 0x56, 0xDB)
C_H2   = RGBColor(0x1D, 0x3F, 0xAF)
C_H3   = RGBColor(0x36, 0x41, 0x50)
C_H4   = RGBColor(0x11, 0x18, 0x26)
C_BODY = RGBColor(0x11, 0x18, 0x26)
C_HDR_FG  = RGBColor(0xFF, 0xFF, 0xFF)
C_CODE_FG = RGBColor(0x1E, 0x25, 0x33)
C_LINK    = RGBColor(0x1D, 0x3F, 0xAF)

# ── XML helpers ───────────────────────────────────────────────────────────────
def _shading(owner_el, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    owner_el.append(shd)

def set_cell_bg(cell, fill_hex):
    _shading(cell._tc.get_or_add_tcPr(), fill_hex)

def set_para_bg(para, fill_hex):
    _shading(para._p.get_or_add_pPr(), fill_hex)

def set_table_borders(table, color="000000", sz=4):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        bdr.append(el)
    tblPr.append(bdr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    '9360')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        run = para.add_run()
        run.font.name = 'Tahoma'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        for field, text in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
            if field == 'instrText':
                el = OxmlElement('w:instrText')
                el.text = text
            else:
                el = OxmlElement('w:fldChar')
                el.set(qn('w:fldCharType'), field)
            run._r.append(el)

def page_break(doc):
    p  = doc.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)

def bottom_border(para, color, sz=6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement('w:bottom')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    str(sz))
    el.set(qn('w:space'), '1')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

# ── run styler ────────────────────────────────────────────────────────────────
def _run(para, text, font, size, bold=False, italic=False, color=None, underline=False):
    r = para.add_run(text)
    r.font.name      = font
    r.font.size      = Pt(size)
    r.font.bold      = bold
    r.font.italic    = italic
    r.font.underline = underline
    if color:
        r.font.color.rgb = color
    return r

# ── inline markdown → runs ────────────────────────────────────────────────────
INLINE = re.compile(
    r'(`+)(.*?)\1'
    r'|\*\*\*(.+?)\*\*\*'
    r'|\*\*(.+?)\*\*'
    r'|__(.+?)__'
    r'|\*(.+?)\*'
    r'|_(.+?)_'
    r'|\[([^\]]+)\]\([^\)]+\)',
    re.DOTALL
)

def inline_runs(para, text, font='Tahoma', size=10.5, color=C_BODY):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _run(para, text[pos:m.start()], font, size, color=color)
        if m.group(1):                              # `code`
            r = para.add_run(m.group(2))
            r.font.name = 'Courier New'
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(0xD6, 0x33, 0x6C)
        elif m.group(3):                            # ***bold+italic***
            _run(para, m.group(3), font, size, bold=True, italic=True, color=color)
        elif m.group(4) or m.group(5):              # **bold**
            _run(para, m.group(4) or m.group(5), font, size, bold=True, color=color)
        elif m.group(6) or m.group(7):              # *italic*
            _run(para, m.group(6) or m.group(7), font, size, italic=True, color=color)
        elif m.group(8):                            # [link text](url)
            _run(para, m.group(8), font, size, color=C_LINK, underline=True)
        pos = m.end()
    if pos < len(text):
        _run(para, text[pos:], font, size, color=color)

# ── paragraph builders ────────────────────────────────────────────────────────
def add_heading(doc, text, level):
    clean = re.sub(r'[`*_]', '', text).strip()
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        _run(p, clean, 'Tahoma', 22, bold=True, color=C_H1)
        pf.space_before = Pt(14); pf.space_after = Pt(6)
        bottom_border(p, '1A56DB', sz=6)
    elif level == 2:
        _run(p, clean, 'Tahoma', 14, bold=True, color=C_H2)
        pf.space_before = Pt(10); pf.space_after = Pt(4)
        bottom_border(p, 'D1D5DB', sz=4)
    elif level == 3:
        _run(p, clean, 'Tahoma', 11.5, bold=True, color=C_H3)
        pf.space_before = Pt(8);  pf.space_after = Pt(3)
    elif level == 4:
        _run(p, clean, 'Tahoma', 10.5, bold=True, color=C_H4)
        pf.space_before = Pt(6);  pf.space_after = Pt(2)
    else:
        _run(p, clean, 'Verdana', 10.5, bold=True, color=C_H4)
        pf.space_before = Pt(5);  pf.space_after = Pt(2)

def add_body(doc, text):
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    inline_runs(p, text)
    return p

def add_code_line(doc, text):
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.3)
    set_para_bg(p, 'FFFFFF')
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8)
    r.font.color.rgb = C_CODE_FG

def add_bullet(doc, text, indent=0):
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.4)
    _run(p, '• ', 'Tahoma', 10.5, bold=True, color=C_BODY)
    inline_runs(p, text)

def add_numbered(doc, text, n):
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    _run(p, f'{n}. ', 'Tahoma', 10.5, bold=True, color=C_BODY)
    inline_runs(p, text)

def add_hr(doc):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    bottom_border(p, 'E5E7EB', sz=6)

def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl   = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl, sz=6)
    C_BLACK = RGBColor(0x00, 0x00, 0x00)
    for ri, row_data in enumerate(rows):
        is_hdr = ri == 0
        for ci, cell_text in enumerate(row_data[:ncols]):
            cell = tbl.rows[ri].cells[ci]
            set_cell_bg(cell, 'FFFFFF')
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            # cell padding
            tcPr  = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side in ('top','left','bottom','right'):
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'),    '80')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            if is_hdr:
                _run(para, cell_text.strip(), 'Tahoma', 9.5, bold=True, color=C_BLACK)
            else:
                inline_runs(para, cell_text.strip(), font='Tahoma', size=9.5, color=C_BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── markdown → docx ───────────────────────────────────────────────────────────
def render(doc, md):
    lines   = md.splitlines()
    i       = 0
    num_ctr = {}

    while i < len(lines):
        line = lines[i]

        # fenced code block
        if line.strip().startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                add_code_line(doc, lines[i])
                i += 1
            i += 1
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # horizontal rule
        if re.match(r'^-{3,}$', line.strip()):
            add_hr(doc)
            i += 1
            continue

        # table
        if line.strip().startswith('|') and '|' in line[1:]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not all(re.match(r'^[-: ]+$', c) for c in cells if c):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        # heading
        m = re.match(r'^(#{1,5})\s+(.+)', line)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # bullet
        m = re.match(r'^(\s*)[-*+]\s+(.+)', line)
        if m:
            add_bullet(doc, m.group(2), indent=len(m.group(1)) // 2)
            i += 1
            continue

        # numbered list
        m = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if m:
            lvl = len(m.group(1)) // 2
            num_ctr[lvl] = num_ctr.get(lvl, 0) + 1
            add_numbered(doc, m.group(2), num_ctr[lvl])
            i += 1
            continue
        else:
            num_ctr = {}

        # blank line
        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # normal paragraph
        add_body(doc, line.strip())
        i += 1

# ── cover ─────────────────────────────────────────────────────────────────────
def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run("Fullstack Deployment using\nDocker, Kubernetes,\nTerraform & Jenkins")
    r.font.name = 'Tahoma'; r.font.size = Pt(36)
    r.font.bold = True;     r.font.color.rgb = C_H1

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(10)
    p2.paragraph_format.space_after  = Pt(40)
    _run(p2,
         "Production-grade infrastructure on AWS — EKS, ALB, RDS, Redis, "
         "ECR, Route53 — deployed and managed entirely through code.",
         'Tahoma', 13, color=RGBColor(0x37, 0x41, 0x51))

    for line in [
        "Deployed at: https://www.ankit.services",
        "AWS Account: 668076964228   |   Region: us-east-1",
    ]:
        pm = doc.add_paragraph()
        pm.paragraph_format.space_before = Pt(2)
        pm.paragraph_format.space_after  = Pt(2)
        _run(pm, line, 'Tahoma', 9.5, color=RGBColor(0x6B, 0x72, 0x80))

    page_break(doc)

# ── main ──────────────────────────────────────────────────────────────────────
def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Tahoma'
    doc.styles['Normal'].font.size = Pt(10.5)

    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    add_cover(doc)
    add_page_number(doc)

    for idx, fname in enumerate(FILES):
        path = os.path.join(DOCS, fname)
        with open(path) as f:
            md = f.read()
        render(doc, md)
        if idx < len(FILES) - 1:
            page_break(doc)

    out = os.path.join(DOCS, "DevOps_Documentation.docx")
    doc.save(out)
    size_kb = os.path.getsize(out) // 1024
    print(f"Done → {out}  ({size_kb} KB)")

if __name__ == "__main__":
    main()
