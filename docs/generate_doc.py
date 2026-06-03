#!/usr/bin/env python3
"""
generate_doc.py — generates the complete engineering documentation .docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_break(doc):
    doc.add_page_break()

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_subheading(doc, text, level=2):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_code_block(doc, code_text):
    """Add a shaded code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFFFF')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x63)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    return p

def add_note_box(doc, text, label="NOTE"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF9E6')
    pPr.append(shd)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB8, 0x6F, 0x00)
    run2 = p.add_run(text)
    run2.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_placeholder(doc, text, detail=""):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F4FD')
    pPr.append(shd)
    run = p.add_run(f"[ {text} ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x6E, 0xC9)
    run.font.size = Pt(10)
    if detail:
        run2 = p.add_run(f"\n{detail}")
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_two_col_table(doc, rows_data, header=None, col_widths=None):
    """rows_data: list of (col1, col2) tuples"""
    n_cols = len(rows_data[0]) if rows_data else 2
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header:
        row = table.add_row()
        for i, h in enumerate(header):
            cell = row.cells[i]
            cell.text = h
            set_cell_bg(cell, 'FFFFFF')
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    return table

# ─────────────────────────────────────────────────────────────
# Document build
# ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Fullstack Application Deployment")
title_run.bold = True
title_run.font.size = Pt(26)
title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run("Using Docker · Kubernetes · Terraform · Jenkins on AWS EKS")
sub_run.bold = True
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x0F, 0x3E, 0x7E)

doc.add_paragraph()
doc.add_paragraph()

cover_table = doc.add_table(rows=6, cols=2)
cover_table.style = 'Table Grid'
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
fields = [
    ("Project", "Todo Fullstack — Production Deployment"),
    ("Prepared by", "Ankit (DevOps / Cloud Engineer)"),
    ("AWS Account", "668076964228  |  Region: us-east-1"),
    ("Domain", "https://www.ankit.services"),
    ("Repository", "github.com/ankitBytes/todo-fullstack"),
    ("Date", "May 2026"),
]
for i, (k, v) in enumerate(fields):
    row = cover_table.rows[i]
    row.cells[0].text = k
    set_cell_bg(row.cells[0], 'FFFFFF')
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = v
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

stack_p = doc.add_paragraph()
stack_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
stack_run = stack_p.add_run("Technology Stack")
stack_run.bold = True
stack_run.font.size = Pt(11)
stack_run.font.color.rgb = RGBColor(0x0F, 0x3E, 0x7E)

tech_stack_text = (
    "Node.js · Express.js · MySQL (RDS) · Redis (ElastiCache) · "
    "Docker · Kubernetes (EKS 1.30) · Terraform · Jenkins · "
    "AWS ALB · AWS ECR · AWS ACM · AWS Secrets Manager · "
    "Route53 · CloudWatch · Cluster Autoscaler · HPA · IRSA"
)
ts_p = doc.add_paragraph()
ts_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ts_run = ts_p.add_run(tech_stack_text)
ts_run.font.size = Pt(9.5)
ts_run.italic = True

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "Table of Contents", level=1)
toc_items = [
    ("1",  "Executive Summary"),
    ("2",  "Project Overview"),
    ("3",  "System Architecture Overview"),
    ("4",  "Complete Technology Stack"),
    ("5",  "Repository & Folder Structure"),
    ("6",  "AWS Services Used"),
    ("7",  "Infrastructure Design"),
    ("8",  "Terraform Architecture"),
    ("9",  "Containerization Strategy"),
    ("10", "Kubernetes Architecture"),
    ("11", "CI/CD Pipeline"),
    ("12", "Application Deployment Flow"),
    ("13", "End-to-End Request Lifecycle"),
    ("14", "Security Implementation"),
    ("15", "Monitoring and Logging"),
    ("16", "Scalability Considerations"),
    ("17", "Cost Optimization"),
    ("18", "Production Best Practices Implemented"),
    ("19", "Challenges Faced and Solutions"),
    ("20", "Key Learnings"),
    ("21", "Future Improvements"),
    ("22", "Deployment Commands Reference"),
    ("23", "Troubleshooting Cheat Sheet"),
    ("24", "Conclusion"),
    ("25", "References"),
    ("26", "Prometheus & Grafana Monitoring Stack"),
    ("27", "Backend Prometheus Metrics Instrumentation"),
]
toc_table = doc.add_table(rows=len(toc_items), cols=2)
toc_table.style = 'Table Grid'
for i, (num, title) in enumerate(toc_items):
    toc_table.rows[i].cells[0].text = num
    toc_table.rows[i].cells[1].text = title
    if i % 2 == 0:
        set_cell_bg(toc_table.rows[i].cells[0], 'FFFFFF')
        set_cell_bg(toc_table.rows[i].cells[1], 'FFFFFF')
    for c in toc_table.rows[i].cells:
        for para in c.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
add_heading(doc, "1. Executive Summary", level=1)
add_para(doc, (
    "This document describes the complete design, implementation, and operational details of a "
    "production-grade fullstack Todo application deployed on Amazon Web Services. The project "
    "demonstrates an end-to-end DevOps workflow — from infrastructure-as-code through containerization, "
    "Kubernetes orchestration, and continuous deployment — following patterns that scale to real "
    "enterprise workloads."
))
add_para(doc, (
    "The application itself is intentionally simple: a Node.js/Express REST API backed by MySQL and "
    "Redis, served by a vanilla-JS single-page frontend. That simplicity is deliberate. The complexity "
    "and learning value are entirely in the infrastructure: multi-tier VPC networking, EKS with managed "
    "node groups, IRSA-based IAM delegation, Secrets Manager integration through the CSI driver, "
    "path-based ALB routing enforced by TLS 1.3, and a fully automated Jenkins pipeline that handles "
    "both application deployments and Terraform infrastructure changes."
))
add_para(doc, (
    "The stack is deployed on AWS Account 668076964228 in us-east-1 and is reachable at "
    "https://www.ankit.services. Jenkins CI/CD is accessible at https://jenkins.ankit.services."
))

doc.add_paragraph()
add_two_col_table(doc, [
    ("Live Application URL",        "https://www.ankit.services"),
    ("Jenkins CI/CD URL",           "https://jenkins.ankit.services"),
    ("EKS Cluster",                 "todo-tf-cluster-dev  |  Kubernetes 1.30"),
    ("AWS Region",                  "us-east-1  (us-east-1a, us-east-1b)"),
    ("Infrastructure as Code",      "Terraform 1.9 / AWS Provider 5.100"),
    ("Container Registry",          "Amazon ECR — todo-frontend, todo-backend"),
    ("Database",                    "RDS MySQL 8.0 on db.t3.micro (20 GiB gp3, encrypted)"),
    ("Cache",                       "ElastiCache Redis 7.1 on cache.t3.micro"),
    ("TLS Certificate",             "AWS ACM — e4a2f397-c129-4501-b3bd-b4ab9d6f22d7"),
    ("Secrets Management",          "AWS Secrets Manager + Secrets Store CSI Driver"),
], header=["Parameter", "Value"])
add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 2 — PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════
add_heading(doc, "2. Project Overview", level=1)

add_subheading(doc, "2.1  Project Goal", level=2)
add_para(doc, (
    "Build and operate a production-ready cloud-native application deployment that demonstrates "
    "the full DevOps lifecycle: infrastructure provisioning, containerization, Kubernetes orchestration, "
    "secure secret management, automated CI/CD, and observability — all on AWS."
))

add_subheading(doc, "2.2  Business Objective", level=2)
add_para(doc, (
    "The primary objective was not the Todo application itself but the surrounding infrastructure "
    "ecosystem: proving that a small engineering team can take code from a git commit to a live, "
    "HTTPS-secured, auto-scaling production deployment with zero manual steps beyond the initial "
    "terraform apply and Jenkins bootstrap."
))

add_subheading(doc, "2.3  Problem Statement", level=2)
add_para(doc, (
    "Deploying modern applications in a cloud-native way requires orchestrating dozens of moving parts — "
    "VPC networking, IAM permissions, container builds, Kubernetes manifests, DNS records, TLS certificates, "
    "secrets management, and a CI/CD pipeline — all in a repeatable, auditable way. The goal was to wire "
    "all of these together with minimal manual intervention and without hard-coding credentials anywhere "
    "in the codebase or pipeline."
))

add_subheading(doc, "2.4  Expected Outcome", level=2)
bullets = [
    "Any git push to the main branch that touches app/frontend/ or app/backend/ triggers a build, ECR push, and zero-downtime rolling deployment automatically.",
    "Any git push that modifies infra/ triggers Terraform plan + apply through Jenkins without manual CLI access.",
    "All database and Redis credentials are injected from AWS Secrets Manager at pod startup — never stored in environment files or git.",
    "The application scales horizontally via HPA (pods) and vertically via Cluster Autoscaler (nodes) without operator intervention.",
    "A second engineer can recreate the entire stack from scratch by running terraform init && terraform apply with the correct tfvars.",
]
for b in bullets:
    add_bullet(doc, b)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 3 — SYSTEM ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════
add_heading(doc, "3. System Architecture Overview", level=1)

add_para(doc, (
    "The architecture follows a classic cloud-native three-tier pattern hardened for production: "
    "a stateless presentation tier (frontend), a stateless application tier (backend API), and "
    "a persistent data tier (RDS + ElastiCache) — all inside a custom VPC with strict subnet isolation."
))

add_placeholder(doc,
    "Insert Overall Architecture Diagram Here",
    "High-level AWS architecture diagram showing VPC, public/private subnets, EKS cluster, ALB, "
    "RDS, ElastiCache, ECR, Route53, ACM, Secrets Manager, and Jenkins. "
    "Reference: docs/diagrams/Overall Architecture-2026-05-19-085235.png"
)

add_subheading(doc, "3.1  Architecture Layers", level=2)
add_two_col_table(doc, [
    ("DNS & TLS",           "Route53 A-record (alias) → ALB  |  ACM certificate for TLS 1.3 termination"),
    ("Ingress Layer",       "Application Load Balancer (internet-facing) — HTTP→HTTPS redirect, path-based routing"),
    ("Orchestration",       "Amazon EKS 1.30 — managed node group (t3.medium × 2, max 3) in private subnets"),
    ("Application",         "Frontend: serve@14 (static files)  |  Backend: Node.js/Express REST API"),
    ("Data",                "RDS MySQL 8.0 (private subnet, encrypted)  |  ElastiCache Redis 7.1 (private subnet)"),
    ("Secrets",             "AWS Secrets Manager + Secrets Store CSI Driver — zero-credential container startup"),
    ("Image Registry",      "Amazon ECR — todo-frontend and todo-backend repositories"),
    ("CI/CD",               "Jenkins (running inside EKS, jenkins namespace) — Kubernetes plugin + DinD agent pods"),
    ("Infrastructure IaC",  "Terraform 1.9 — 13 modules, S3 remote state, native S3 locking"),
    ("Observability",       "CloudWatch log groups for EKS control plane, application, backend, and frontend"),
], header=["Layer", "Components"])

add_subheading(doc, "3.2  Key Design Principles", level=2)
principles = [
    "Private-by-default: EKS nodes, RDS, and ElastiCache all live in private subnets. Only the ALB is internet-facing.",
    "Zero long-lived credentials: Jenkins and backend pods use IRSA (IAM Roles for Service Accounts) — no AWS access keys stored anywhere.",
    "Change-gated pipeline: Terraform runs only when infra/ changes; image builds run only when the respective app directory changes. This prevents accidental infrastructure churn on pure code pushes.",
    "Automated rollback: kubectl rollout status with a 180-second timeout triggers kubectl rollout undo if any deployment fails to become healthy.",
    "Immutable tags with latest alias: Every CI build produces a versioned tag (v<BUILD_NUMBER>). The :latest tag is also pushed so K8s manifests can reference it without tag-patching in most scenarios.",
]
for p in principles:
    add_bullet(doc, p)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 4 — COMPLETE TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════
add_heading(doc, "4. Complete Technology Stack", level=1)

add_two_col_table(doc, [
    ("Runtime",              "Node.js 20 LTS (Alpine base image)"),
    ("Web Framework",        "Express.js 4.18"),
    ("Database Driver",      "mysql2 3.6 — promise-based connection pooling (limit 10)"),
    ("Cache Client",         "ioredis 5.3 — lazy-connect, error-tolerant"),
    ("CORS",                 "cors 2.8 — configured to allow only https://www.ankit.services"),
    ("Configuration",        "dotenv 16 — picks up .env in development; secrets injected via CSI in K8s"),
    ("Frontend",             "Vanilla HTML5 / CSS3 / JavaScript — no build step, served by serve@14"),
    ("Containerization",     "Docker — multi-layer non-root images on node:20-alpine base"),
    ("Dev Compose",          "Docker Compose — nginx reverse proxy + mysql + redis + backend"),
    ("Orchestration",        "Kubernetes 1.30 on Amazon EKS (managed node groups)"),
    ("IaC",                  "Terraform 1.9 — AWS Provider 5.100, Helm Provider 2.17, TLS Provider 4.3"),
    ("CI/CD",                "Jenkins (Helm chart) — Kubernetes plugin, DinD sidecar, IRSA auth"),
    ("Load Balancing",       "AWS Application Load Balancer + AWS Load Balancer Controller (IRSA)"),
    ("DNS",                  "Amazon Route53 — A-record alias to ALB (app + jenkins subdomains)"),
    ("TLS",                  "AWS ACM — wildcard or SAN cert; TLS policy ELBSecurityPolicy-TLS13-1-2-2021-06"),
    ("Image Registry",       "Amazon ECR — lifecycle policy retains 5 tagged images"),
    ("Secrets",              "AWS Secrets Manager + Secrets Store CSI Driver 1.4.4 + AWS Provider 0.3.9"),
    ("Autoscaling",          "HPA v2 (CPU 70% / Memory 80%) + Cluster Autoscaler (ASG target-tracking 60% CPU)"),
    ("Monitoring",           "CloudWatch — 14-day log retention; EKS control plane full audit logging"),
    ("State Backend",        "S3 bucket todo-tf-state-668076964228 with versioning + native S3 lock file"),
], header=["Component", "Details"])

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 5 — REPOSITORY & FOLDER STRUCTURE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "5. Repository & Folder Structure", level=1)

add_para(doc, (
    "The repository is organized into three top-level directories — app/, infra/, and docs/ — "
    "keeping application code, infrastructure code, and documentation cleanly separated. "
    "This separation matters for the CI/CD pipeline: the Jenkinsfile uses changeset conditions "
    "to run builds only when the relevant subtree has changed."
))

add_code_block(doc, """\
todo-fullstack/
├── app/
│   ├── backend/
│   │   ├── src/
│   │   │   ├── app.js                  # Express server entrypoint
│   │   │   ├── db.js                   # MySQL connection pool with retry
│   │   │   ├── redis.js                # ioredis client setup
│   │   │   ├── controllers/
│   │   │   │   └── todo.controller.js  # CRUD handlers + Redis cache-aside
│   │   │   └── routes/
│   │   │       └── todo.routes.js      # Route registration
│   │   └── package.json
│   ├── frontend/
│   │   ├── index.html
│   │   ├── app.js                      # Vanilla JS SPA
│   │   └── style.css
│   ├── docker/
│   │   ├── Dockerfile.backend          # node:20-alpine, non-root, npm ci --omit=dev
│   │   ├── Dockerfile.frontend         # node:20-alpine, serve@14, non-root
│   │   ├── Dockerfile.jenkins-agent    # inbound-agent + Docker CLI + AWS CLI + kubectl + Terraform
│   │   ├── docker-compose.yaml         # Local dev stack (nginx + backend + mysql + redis)
│   │   ├── nginx.dev.conf              # Mirrors ALB path routing for local testing
│   │   └── .dockerignore
│   ├── k8s/
│   │   ├── namespace.yaml              # todo namespace
│   │   ├── deployment.yaml             # frontend + backend Deployments (2 replicas, anti-affinity)
│   │   ├── service.yaml                # ClusterIP services (port 80 → 3000)
│   │   ├── ingress.yaml                # ALB Ingress — path-based routing, TLS, SSL redirect
│   │   ├── configmap.yaml              # Non-sensitive config (PORT, DB_NAME, ALLOWED_ORIGIN)
│   │   ├── secret.yaml                 # Kubernetes Secret (populated by CSI driver)
│   │   ├── secretproviderclass.yaml    # Secrets Store CSI → AWS Secrets Manager mapping
│   │   ├── serviceaccount.yaml         # IRSA annotation for todo-backend-sa
│   │   ├── hpa.yaml                    # HPA v2 for frontend (2-6) and backend (2-10)
│   │   ├── pdb.yaml                    # PodDisruptionBudget minAvailable:1 for each service
│   │   ├── networkpolicy.yaml          # Ingress/egress rules per pod label
│   │   └── targetgroupbinding.yaml     # TargetGroupBinding CRD — bind TGs to K8s services
│   └── ci/
│       └── deploy.sh                   # Manual deployment helper with preflight + rollback
├── infra/
│   ├── ci-cd/
│   │   ├── Jenkinsfile                 # Declarative pipeline — parallel app + sequential Terraform
│   │   └── jenkins-values.yaml         # Helm values for Jenkins — DinD, IRSA, ALB ingress
│   └── terraform/
│       ├── main.tf                     # Root module — wires all modules together
│       ├── variables.tf                # All input variables with defaults
│       ├── outputs.tf                  # Key outputs (ARNs, endpoints, URLs)
│       ├── providers.tf                # AWS + Helm + TLS providers; cluster auth
│       ├── backend.tf                  # S3 remote state + locking
│       ├── versions.tf                 # Provider version constraints
│       ├── jenkins-irsa.tf             # Jenkins IRSA role + ECR/EKS/PowerUser policies
│       ├── environments/
│       │   ├── dev/terraform.tfvars    # Dev environment values
│       │   └── prod/terraform.tfvars   # Production environment values
│       └── modules/
│           ├── vpc/                    # VPC, subnets, IGW, NAT GW, route tables
│           ├── security-groups/        # SGs for EKS nodes, RDS, Redis, ALB
│           ├── eks/                    # EKS cluster, node group, OIDC, IRSA for EBS CSI + ALB controller
│           ├── ecr/                    # ECR repositories with lifecycle policy
│           ├── rds/                    # RDS MySQL 8.0 with encrypted gp3 storage
│           ├── redis/                  # ElastiCache Redis 7.1 replication group
│           ├── alb/                    # ALB + target groups + listeners + path rules
│           ├── route53/                # A-record alias for app + jenkins subdomains
│           ├── secrets-manager/        # Secrets Manager secret + IRSA role for backend
│           ├── cloudwatch/             # CloudWatch log groups
│           ├── s3/                     # S3 bucket for assets
│           ├── asg/                    # ASG target-tracking scaling policy (CPU 60%)
│           └── autoscalling/           # Cluster Autoscaler IRSA role + policy
└── docs/
    ├── diagrams/                       # Architecture diagrams (.png)
    └── *.md                            # Supporting documentation
""")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 6 — AWS SERVICES USED
# ══════════════════════════════════════════════════════════════
add_heading(doc, "6. AWS Services Used", level=1)

services = [
    (
        "Amazon EKS (Elastic Kubernetes Service)",
        "Core orchestration platform. Runs Kubernetes 1.30 with a managed node group of t3.medium instances "
        "across two AZs. The control plane is fully managed; nodes live in private subnets. OIDC provider "
        "enables IRSA for passwordless IAM delegation to pods."
    ),
    (
        "Amazon ECR (Elastic Container Registry)",
        "Private Docker registry for todo-frontend and todo-backend images. Jenkins pushes versioned "
        "(v<BUILD_NUMBER>) and :latest tags after each successful build. A lifecycle policy keeps the last "
        "5 tagged images, preventing unbounded storage growth."
    ),
    (
        "Amazon RDS (MySQL 8.0)",
        "Relational database in a private subnet with gp3 encrypted storage. 7-day automated backups "
        "(3 days in dev), maintenance window Mon 04:00-05:00 UTC, slow query and error logs forwarded "
        "to CloudWatch. Multi-AZ is disabled in dev to save cost but variable-ready for prod."
    ),
    (
        "Amazon ElastiCache (Redis 7.1)",
        "In-memory cache used for a 60-second cache-aside TTL on the GET /api/todos endpoint. "
        "Placed in private subnets with at-rest encryption. Single-node in dev (automatic failover "
        "disabled); the variable allows switching to multi-node for prod."
    ),
    (
        "AWS Application Load Balancer (ALB)",
        "Internet-facing L7 load balancer managed by the AWS Load Balancer Controller inside EKS. "
        "Terminates TLS with the ACM cert, enforces HTTP→HTTPS 301 redirect, and routes /api/* and "
        "/health to the backend target group while / routes to the frontend target group."
    ),
    (
        "AWS Secrets Manager",
        "Stores all sensitive config (DB host, user, password, Redis host/port) as a single JSON secret "
        "named todo-dev-app-credentials. The Secrets Store CSI Driver mounts these values directly "
        "into the backend pod filesystem and syncs them as a Kubernetes Secret. No secrets in git."
    ),
    (
        "Amazon Route53",
        "Hosts the ankit.services public zone. Two A-record aliases: www.ankit.services → app ALB and "
        "jenkins.ankit.services → Jenkins ALB. Evaluate-target-health is enabled so DNS fails over if "
        "the ALB is unhealthy."
    ),
    (
        "AWS ACM (Certificate Manager)",
        "Provides the TLS certificate (ARN: e4a2f397) attached to both ALBs. The certificate covers "
        "both www.ankit.services and jenkins.ankit.services. TLS policy ELBSecurityPolicy-TLS13-1-2-2021-06 "
        "enforces TLS 1.2 minimum, preferring TLS 1.3."
    ),
    (
        "Amazon S3",
        "Two purposes: (1) Terraform remote state storage (todo-tf-state-668076964228) with versioning "
        "and native S3 lock file; (2) application asset bucket (todo-assets-dev-668076964228) for future "
        "static content offloading."
    ),
    (
        "Amazon CloudWatch",
        "Four log groups pre-created by Terraform before the EKS cluster starts — this prevents "
        "the race condition where EKS tries to write logs before the group exists. Logs: "
        "/aws/eks/todo-tf-cluster-dev/cluster (EKS control plane), /todo/dev/application, "
        "/todo/dev/backend, /todo/dev/frontend. Retention: 14 days in dev."
    ),
    (
        "AWS IAM (Identity and Access Management)",
        "Every component uses least-privilege IAM. EKS cluster role: AmazonEKSClusterPolicy. "
        "Node role: WorkerNodePolicy + CNI + ECRReadOnly + CloudWatchAgent. "
        "IRSA roles are created for: EBS CSI driver, ALB controller, Cluster Autoscaler, backend pod "
        "(Secrets Manager read), Jenkins (ECR full + EKS describe + PowerUser + IAMFull for Terraform)."
    ),
    (
        "AWS Auto Scaling (ASG)",
        "Cluster Autoscaler watches the EKS node group ASG (min 1, max 3 in dev). Target-tracking "
        "policy triggers scale-out at 60% average CPU. The ASG module looks up the node group ASG "
        "by EKS cluster and node group name tags, keeping it decoupled from the ASG ID."
    ),
]

for name, desc in services:
    add_subheading(doc, name, level=2)
    add_para(doc, desc)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 7 — INFRASTRUCTURE DESIGN
# ══════════════════════════════════════════════════════════════
add_heading(doc, "7. Infrastructure Design", level=1)

add_subheading(doc, "7.1  VPC and Subnets", level=2)
add_para(doc, (
    "A single VPC (10.0.0.0/16) spans two availability zones (us-east-1a, us-east-1b). "
    "The subnet layout follows a standard two-tier design: public subnets host the NAT Gateway "
    "and ALB ENIs; private subnets host everything else."
))

add_placeholder(doc,
    "Insert VPC Network Layout Diagram Here",
    "Shows VPC 10.0.0.0/16 with public subnets (10.0.1.0/24, 10.0.2.0/24) and "
    "private subnets (10.0.10.0/24, 10.0.11.0/24) across us-east-1a and us-east-1b. "
    "Reference: docs/diagrams/VPC Network Layout-2026-05-19-085304.png"
)

add_two_col_table(doc, [
    ("VPC CIDR",              "10.0.0.0/16"),
    ("Public Subnet 1a",      "10.0.1.0/24  (us-east-1a)  — ALB ENIs, NAT GW"),
    ("Public Subnet 1b",      "10.0.2.0/24  (us-east-1b)  — ALB ENIs"),
    ("Private Subnet 1a",     "10.0.10.0/24 (us-east-1a)  — EKS nodes, RDS, Redis"),
    ("Private Subnet 1b",     "10.0.11.0/24 (us-east-1b)  — EKS nodes"),
    ("Internet Gateway",      "Attached to VPC — public subnets route 0.0.0.0/0 → IGW"),
    ("NAT Gateway",           "Single NAT GW in public-1a with Elastic IP — private subnets route 0.0.0.0/0 → NAT"),
    ("Public Route Table",    "0.0.0.0/0 → IGW"),
    ("Private Route Table",   "0.0.0.0/0 → NAT GW"),
], header=["Component", "Detail"])

add_note_box(doc,
    "A single NAT Gateway is a cost tradeoff. For prod, deploy a NAT GW per AZ to eliminate "
    "cross-AZ data transfer charges and AZ-level single points of failure. The variable "
    "availability_zones makes this extension straightforward.",
    "ARCH NOTE"
)

add_subheading(doc, "7.2  Security Groups", level=2)
add_para(doc, (
    "Security groups are defined as a dedicated Terraform module so changes to one SG don't "
    "force replacement of the EKS node group. The key security group rules are:"
))
add_two_col_table(doc, [
    ("ALB SG",           "Inbound: TCP 80, 443 from 0.0.0.0/0  |  Outbound: all"),
    ("EKS Node SG",      "Inbound: node-to-node (self), port 3000 from ALB SG, kubelet/HTTPS/ephemeral from cluster SG  |  Outbound: all"),
    ("RDS SG",           "Inbound: TCP 3306 from EKS Node SG only  |  Outbound: all"),
    ("Redis SG",         "Inbound: TCP 6379 from EKS Node SG only  |  Outbound: all"),
], header=["Security Group", "Rules"])

add_subheading(doc, "7.3  IAM Strategy (IRSA)", level=2)
add_para(doc, (
    "All IAM permissions follow the IAM Roles for Service Accounts (IRSA) pattern. This eliminates "
    "the need for AWS_ACCESS_KEY_ID / AWS_SECRET_ACCESS_KEY in any pod or pipeline. IRSA works "
    "by federating pod identity through the EKS OIDC provider: a service account annotated with "
    "an IAM role ARN allows the pod to exchange a Kubernetes projected service account token for "
    "temporary AWS credentials via STS AssumeRoleWithWebIdentity."
))

add_placeholder(doc,
    "Insert IRSA Flow Diagram Here",
    "Shows: Pod ServiceAccount → OIDC Provider → STS AssumeRoleWithWebIdentity → IAM Role → AWS Service. "
    "Reference: docs/diagrams/IRSA Flow-2026-05-19-085411.png"
)

add_two_col_table(doc, [
    ("todo-backend-sa",           "todo-dev-todo-backend-role  →  SecretsManager:GetSecretValue on todo-dev-app-credentials"),
    ("aws-load-balancer-controller", "todo-dev-alb-controller-role  →  ALB controller policy (create/manage ALBs, TGs, listeners)"),
    ("ebs-csi-controller-sa",     "todo-dev-ebs-csi-role  →  AmazonEBSCSIDriverPolicy"),
    ("cluster-autoscaler",        "todo-dev-cluster-autoscaler-role  →  ASG describe/scale, EC2 launch template describe"),
    ("jenkins (jenkins ns)",       "todo-dev-jenkins-irsa-role  →  ECRFullAccess + EKSClusterPolicy + PowerUserAccess + IAMFullAccess"),
], header=["Kubernetes ServiceAccount", "IAM Role & Permissions"])

add_subheading(doc, "7.4  DNS and TLS", level=2)
add_para(doc, (
    "Route53 hosts the ankit.services zone. Two alias A-records are managed by Terraform: "
    "www.ankit.services → app ALB and jenkins.ankit.services → Jenkins ALB. "
    "Using alias records (rather than CNAME) avoids the DNS anti-pattern of a CNAME at zone apex "
    "and enables health-check integration through evaluate_target_health = true."
))
add_para(doc, (
    "TLS is terminated at the ALB using an ACM certificate. The Ingress annotation "
    "alb.ingress.kubernetes.io/ssl-policy: ELBSecurityPolicy-TLS13-1-2-2021-06 enforces "
    "TLS 1.2 minimum, prefers TLS 1.3 cipher suites, and disables deprecated TLS 1.0/1.1 "
    "connections at the load balancer level."
))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 8 — TERRAFORM ARCHITECTURE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "8. Terraform Architecture", level=1)

add_subheading(doc, "8.1  Module Structure", level=2)
add_para(doc, (
    "The Terraform codebase uses a flat root module that composes 13 child modules. "
    "Each module owns a single AWS resource domain (VPC, EKS, RDS, etc.) with its own "
    "variables, outputs, and no inter-module cross-dependencies except through root-level wiring. "
    "This keeps modules reusable in isolation."
))

add_two_col_table(doc, [
    ("vpc",               "VPC, public/private subnets, IGW, NAT GW, route tables, associations"),
    ("security-groups",   "All security groups and cross-SG rules"),
    ("eks",               "EKS cluster, OIDC provider, launch template, node group, EKS addons, IRSA roles"),
    ("ecr",               "ECR repository with lifecycle policy — instantiated twice (frontend, backend)"),
    ("rds",               "RDS MySQL 8.0 instance, subnet group, parameter group"),
    ("redis",             "ElastiCache replication group, subnet group"),
    ("alb",               "ALB, target groups (ip type), listeners, path-based listener rule"),
    ("route53",           "A-record aliases for app and jenkins subdomains"),
    ("secrets-manager",   "Secrets Manager secret + version, IRSA role for backend pod"),
    ("cloudwatch",        "CloudWatch log groups (EKS cluster + app + backend + frontend)"),
    ("s3",                "S3 bucket for application assets"),
    ("asg",               "ASG target-tracking scaling policy on the EKS node group"),
    ("autoscalling",      "Cluster Autoscaler IRSA role and policy"),
], header=["Module", "Resources Created"])

add_subheading(doc, "8.2  Remote State", level=2)
add_code_block(doc, """\
# S3 backend — one-time bootstrap before first terraform init:
aws s3 mb s3://todo-tf-state-668076964228 --region us-east-1
aws s3api put-bucket-versioning \\
    --bucket todo-tf-state-668076964228 \\
    --versioning-configuration Status=Enabled

# backend.tf (root):
terraform {
  backend "s3" {
    bucket       = "todo-tf-state-668076964228"
    key          = "todo-fullstack/default.tfstate"
    region       = "us-east-1"
    encrypt      = true
    use_lockfile = true   # native S3 locking (no DynamoDB required — Terraform 1.10+)
  }
}""")

add_para(doc, (
    "State is encrypted at rest (S3 SSE) and versioned so previous states can be recovered. "
    "Native S3 locking (use_lockfile = true) was used instead of DynamoDB to reduce the number "
    "of resources needed for the bootstrap. This requires Terraform 1.10+ on the client side."
))

add_subheading(doc, "8.3  Provider Configuration", level=2)
add_para(doc, (
    "Three providers are configured at root level. The Helm provider is initialized after the EKS "
    "cluster is available by reading cluster endpoint and CA cert from data sources "
    "(aws_eks_cluster + aws_eks_cluster_auth). This avoids chicken-and-egg ordering problems "
    "that arise when trying to use depends_on with providers."
))
add_para(doc, "Global tags (Project, Environment, ManagedBy=Terraform) are applied via provider default_tags "
              "rather than repeating tag blocks in every resource.")

add_subheading(doc, "8.4  In-cluster Helm Releases", level=2)
add_para(doc, (
    "Two Helm releases are managed by Terraform (not manually applied): "
    "Secrets Store CSI Driver 1.4.4 (with syncSecret and secretRotation enabled) and "
    "AWS Secrets Manager Provider 0.3.9. These are installed into kube-system after the EKS "
    "cluster is ready. Deploying them via Terraform ensures they exist before any application "
    "pods reference the SecretProviderClass CRD."
))

add_subheading(doc, "8.5  Environments", level=2)
add_code_block(doc, """\
# To apply the dev environment:
terraform -chdir=infra/terraform init \\
    -backend-config=environments/dev/backend.tf
terraform -chdir=infra/terraform plan \\
    -var-file=environments/dev/terraform.tfvars \\
    -out=tfplan
terraform -chdir=infra/terraform apply tfplan""")

add_note_box(doc,
    "The rds_password variable is marked sensitive = true and should never be placed in terraform.tfvars. "
    "Set it via: export TF_VAR_rds_password='<password>' before running plan/apply. "
    "In Jenkins, it is injected as a credential binding.",
    "SECURITY"
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 9 — CONTAINERIZATION STRATEGY
# ══════════════════════════════════════════════════════════════
add_heading(doc, "9. Containerization Strategy", level=1)

add_subheading(doc, "9.1  Backend Image (Dockerfile.backend)", level=2)
add_para(doc, (
    "The backend uses node:20-alpine as the base — Alpine reduces the final image size significantly "
    "versus the full Debian node image. Layer-caching is intentional: package.json is copied and "
    "npm ci --omit=dev runs before copying the source files. This means rebuilds that only change "
    "source code skip the npm install layer, making CI builds faster."
))
add_code_block(doc, """\
FROM node:20-alpine
RUN addgroup -S appgroup && adduser -S appuser -G appgroup
WORKDIR /app
COPY package*.json ./
RUN npm ci --omit=dev          # production deps only
COPY src/ ./src/
RUN chown -R appuser:appgroup /app
USER appuser                   # drop root before container starts
EXPOSE 3000
CMD ["node", "src/app.js"]""")

add_subheading(doc, "9.2  Frontend Image (Dockerfile.frontend)", level=2)
add_para(doc, (
    "The frontend has no build step — it is pure HTML/CSS/JS. The serve package (version 14) "
    "is installed globally and serves static files on port 3000. This matches the backend port, "
    "simplifying service configuration. Both images use the same non-root user pattern "
    "(appuser:appgroup) and drop all Linux capabilities in the pod securityContext."
))

add_subheading(doc, "9.3  Jenkins Agent Image (Dockerfile.jenkins-agent)", level=2)
add_para(doc, (
    "The Jenkins agent image is built on jenkins/inbound-agent:latest-jdk17 and layers in "
    "four additional tools required by the Jenkinsfile:"
))
add_two_col_table(doc, [
    ("Docker CLI",   "Connects to the DinD sidecar (tcp://localhost:2375) for docker build/push"),
    ("AWS CLI v2",   "ECR authentication, EKS kubeconfig update, AWS credential resolution"),
    ("kubectl 1.30", "Matches EKS cluster minor version — running a mismatched minor can cause subtle API differences"),
    ("Terraform 1.9","Infrastructure changes triggered by the pipeline"),
], header=["Tool", "Purpose"])

add_subheading(doc, "9.4  Docker Compose (Local Dev)", level=2)
add_para(doc, (
    "The docker-compose.yaml provides a local environment that mirrors production closely: "
    "nginx acts as the reverse proxy (mirroring ALB path routing), the backend connects to a "
    "real MySQL 8.0 container, and Redis 7 is available. Healthcheck conditions ensure the "
    "backend only starts once MySQL is accepting connections — this prevents the startup race "
    "condition where the backend would crash-loop waiting for the DB."
))
add_para(doc, (
    "The nginx.dev.conf routes /api → backend:3000 and / → static files, exactly mirroring "
    "the ALB listener rules. Any routing issue caught in local compose will manifest the same "
    "way in Kubernetes, making debugging consistent across environments."
))

add_subheading(doc, "9.5  Security Hardening in Container Spec", level=2)
bullets = [
    "runAsNonRoot: true + runAsUser: 1000 — enforced at pod level",
    "allowPrivilegeEscalation: false — prevents sudo/setuid privilege escalation",
    "capabilities.drop: [\"ALL\"] — removes all Linux kernel capabilities from the container process",
    "IMDSv2 enforced on nodes: http_tokens = required, hop_limit = 1 — prevents SSRF-based metadata exfiltration from pods",
    "EBS volumes on nodes are encrypted (gp3, 20GB, delete_on_termination = true)",
]
for b in bullets:
    add_bullet(doc, b)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 10 — KUBERNETES ARCHITECTURE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "10. Kubernetes Architecture", level=1)

add_placeholder(doc,
    "Insert EKS Namespace Structure Diagram Here",
    "Shows namespaces: kube-system (ALB controller, CSI driver, CoreDNS, kube-proxy, cluster-autoscaler), "
    "todo (frontend + backend deployments, services, ingress, configmap, secrets), "
    "jenkins (Jenkins controller + agent pods). "
    "Reference: docs/diagrams/EKS Namespace Structure-2026-05-19-085451.png"
)

add_subheading(doc, "10.1  Namespaces", level=2)
add_two_col_table(doc, [
    ("kube-system",  "EKS system components: CoreDNS, kube-proxy, VPC CNI, EBS CSI driver, ALB controller, Secrets Store CSI, Cluster Autoscaler"),
    ("todo",         "Application workloads: frontend deployment, backend deployment, services, ingress, configmap, secret, HPA, PDB, network policies, service account"),
    ("jenkins",      "Jenkins controller pod (persistent 20GB gp2 volume) and ephemeral agent pods"),
], header=["Namespace", "Contents"])

add_subheading(doc, "10.2  Deployments", level=2)
add_para(doc, "Both frontend and backend are configured identically in terms of Kubernetes patterns:")
add_two_col_table(doc, [
    ("Replicas",              "2 (minimum) — ensures availability during node maintenance"),
    ("Rolling Update",        "maxSurge: 1, maxUnavailable: 0 — zero-downtime deploy"),
    ("Pod Anti-Affinity",     "preferredDuringSchedulingIgnoredDuringExecution by topology.kubernetes.io/zone — pods spread across AZs where possible"),
    ("Resource Requests",     "Frontend: 50m CPU / 64Mi RAM  |  Backend: 100m CPU / 128Mi RAM"),
    ("Resource Limits",       "Frontend: 200m CPU / 128Mi RAM  |  Backend: 300m CPU / 256Mi RAM"),
    ("Liveness Probe",        "HTTP GET /health (backend) or / (frontend) — kicks in after 10s, 3 failures to restart"),
    ("Readiness Probe",       "Same paths, earlier start (5s) — pod excluded from load balancing until healthy"),
    ("terminationGrace",      "30 seconds — allows in-flight requests to complete before SIGKILL"),
], header=["Setting", "Value / Notes"])

add_placeholder(doc,
    "Insert Kubernetes Rolling Update Diagram Here",
    "Illustrates the rolling update sequence: old pod running → new pod created (maxSurge:1) → "
    "new pod passes readiness → old pod terminated (maxUnavailable:0). "
    "Reference: docs/diagrams/Kubernetes Rolling Update-2026-05-19-085612.png"
)

add_subheading(doc, "10.3  Services and Ingress", level=2)
add_para(doc, (
    "Both frontend and backend expose ClusterIP services on port 80 → containerPort 3000. "
    "ClusterIP was chosen over NodePort/LoadBalancer because the ALB (managed by the "
    "Load Balancer Controller) routes directly to pod IPs (target-type: ip) — it doesn't need "
    "node-level NodePort exposure."
))
add_para(doc, (
    "The Ingress resource uses the alb IngressClass. The ALB controller reads the annotations "
    "and creates (or reuses) an ALB with the specified cert ARN, listener configuration, and "
    "health check settings. The TargetGroupBinding CRDs explicitly bind the pre-created Terraform "
    "target groups to the Kubernetes services, giving Terraform full control over the ALB lifecycle "
    "while Kubernetes controls target registration."
))

add_subheading(doc, "10.4  Secrets Management (CSI Driver)", level=2)
add_para(doc, (
    "The SecretProviderClass named todo-app-secrets maps the AWS Secrets Manager secret "
    "todo-dev-app-credentials to individual keys using jmesPath selectors. When the backend "
    "pod starts, the CSI driver:"
))
steps = [
    "Authenticates to AWS Secrets Manager using the IRSA token from todo-backend-sa",
    "Retrieves the JSON secret and extracts each field (DB_HOST, DB_USER, DB_PASSWORD, REDIS_HOST, etc.)",
    "Mounts the values as files under /mnt/secrets-store/ in the container",
    "Syncs them into a Kubernetes Secret (todo-secret) in the todo namespace",
    "The Deployment then consumes todo-secret via secretKeyRef env vars",
]
for i, s in enumerate(steps, 1):
    add_bullet(doc, f"{i}. {s}")

add_subheading(doc, "10.5  Horizontal Pod Autoscaler (HPA)", level=2)
add_two_col_table(doc, [
    ("Backend HPA",   "Min 2, Max 10 replicas  |  Scale up: +2 pods/60s when CPU >70% or Memory >80%  |  Scale down: -1 pod/60s after 300s window"),
    ("Frontend HPA",  "Min 2, Max 6 replicas  |  Scale up: +1 pod/60s  |  Scale down: -1 pod/60s after 300s window"),
], header=["HPA", "Configuration"])
add_para(doc, (
    "The scale-down stabilization window (300s = 5 min) prevents thrashing on traffic spikes that "
    "resolve quickly. Scale-up is intentionally aggressive (stabilizationWindowSeconds: 60) "
    "to respond fast to load."
))

add_subheading(doc, "10.6  Pod Disruption Budgets", level=2)
add_para(doc, (
    "PodDisruptionBudgets with minAvailable: 1 are defined for both frontend and backend. "
    "This prevents the Cluster Autoscaler or a manual kubectl drain from evicting all replicas "
    "simultaneously, maintaining availability during node maintenance or scale-down events."
))

add_subheading(doc, "10.7  Network Policies", level=2)
add_para(doc, (
    "Kubernetes NetworkPolicies restrict pod-level traffic. The backend policy allows ingress "
    "only from 10.0.0.0/8 (ALB ENIs in VPC subnets) and from frontend pods, on port 3000. "
    "The frontend policy allows ingress from the same VPC CIDR and allows egress only to backend "
    "pods on port 3000. All other pod-to-pod traffic is denied by default when a NetworkPolicy "
    "selects a pod."
))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 11 — CI/CD PIPELINE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "11. CI/CD Pipeline", level=1)

add_placeholder(doc,
    "Insert End-to-End CI/CD Flow Diagram Here",
    "Shows: GitHub push → GitHub webhook → Jenkins → changeset conditions → parallel frontend/backend "
    "build stages → ECR push → kubectl rolling deploy → rollout status check → rollback on failure. "
    "Also shows the Terraform branch for infra/ changes. "
    "Reference: docs/diagrams/End-to-End CICD Flow-2026-05-19-085645.png"
)

add_subheading(doc, "11.1  Jenkins Deployment", level=2)
add_para(doc, (
    "Jenkins is deployed into the jenkins namespace via Helm (jenkins/jenkins chart). "
    "The controller pod uses a 20GB gp2 persistent volume for job history and plugin storage. "
    "numExecutors is set to 0 on the controller — no builds run there. All job execution happens "
    "in ephemeral Kubernetes agent pods that are created on demand, run the build, and are "
    "deleted when the pipeline finishes."
))
add_para(doc, (
    "Agent pods use a custom image (todo/jenkins-agent ECR) based on jenkins/inbound-agent with "
    "Docker CLI, AWS CLI, kubectl, and Terraform pre-installed. A DinD (Docker-in-Docker) sidecar "
    "container provides the Docker daemon. The agent communicates with the DinD daemon via "
    "DOCKER_HOST=tcp://localhost:2375."
))
add_para(doc, (
    "The Jenkins ServiceAccount is annotated with the jenkins-irsa-role ARN, giving the pipeline "
    "passwordless access to ECR, EKS, and Terraform AWS operations — no static credentials needed."
))

add_subheading(doc, "11.2  Pipeline Stages", level=2)
add_two_col_table(doc, [
    ("Checkout",            "Clones the repository from source control"),
    ("App (parallel)",      "Frontend and backend stages run in parallel. Each has three sub-stages: Build, Push, Deploy — all gated by changeset 'app/frontend/**' or 'app/backend/**'"),
    ("Build Frontend",      "docker build -f app/docker/Dockerfile.frontend -t ECR_FRONTEND:v<BUILD_NUMBER> app/frontend/"),
    ("Push Frontend",       "ECR login via aws ecr get-login-password → docker push versioned + :latest tags"),
    ("Deploy Frontend",     "aws eks update-kubeconfig → kubectl set image → kubectl rollout status (180s timeout) || rollback"),
    ("Build Backend",       "docker build -f app/docker/Dockerfile.backend -t ECR_BACKEND:v<BUILD_NUMBER> app/backend/"),
    ("Push Backend",        "Same ECR login + push pattern as frontend"),
    ("Deploy Backend",      "Same kubectl rolling deploy pattern as frontend"),
    ("Terraform Init",      "Only when changeset 'infra/**' — terraform init -input=false"),
    ("Terraform Validate",  "terraform validate"),
    ("Terraform Plan",      "terraform plan -var-file=environments/dev/terraform.tfvars -out=tfplan"),
    ("Terraform Apply",     "terraform apply -input=false tfplan"),
    ("Cleanup",             "Remove local Docker images and tfplan file"),
], header=["Stage", "Action"])

add_subheading(doc, "11.3  Changeset-Based Stage Gating", level=2)
add_para(doc, (
    "The when { changeset '...' } condition is one of the most important design decisions in "
    "the pipeline. Without it, every commit — even a docs-only change — would trigger all "
    "build, push, and deploy stages. With changeset gating:"
))
bullets = [
    "A commit that only modifies app/backend/ triggers backend build+push+deploy but skips frontend and Terraform.",
    "A commit that only modifies infra/ triggers Terraform stages but skips all Docker builds.",
    "A commit that touches both app/frontend/ and infra/ runs both sets of stages.",
    "This prevents unnecessary ECR image builds, reduces pipeline time, and avoids accidental Terraform applies on pure code pushes.",
]
for b in bullets:
    add_bullet(doc, b)

add_subheading(doc, "11.4  Rollback Strategy", level=2)
add_para(doc, (
    "Every deployment stage uses kubectl rollout status with a 180-second timeout. "
    "If the rollout doesn't complete (new pods don't pass readiness probes in time), "
    "the pipeline immediately executes kubectl rollout undo and exits with a non-zero code. "
    "Kubernetes keeps the previous ReplicaSet around after a rolling update, so the undo "
    "operation is fast and doesn't re-pull images."
))
add_code_block(doc, """\
kubectl rollout status deployment/todo-backend -n todo --timeout=180s \\
  || (kubectl rollout undo deployment/todo-backend -n todo && exit 1)""")

add_subheading(doc, "11.5  Webhook Integration", level=2)
add_para(doc, (
    "GitHub webhooks trigger the Jenkins pipeline on every push to the main branch. "
    "The webhook URL points to https://jenkins.ankit.services/github-webhook/. "
    "The github plugin in Jenkins verifies the webhook payload and triggers the "
    "appropriate multibranch pipeline or freestyle job depending on the configuration."
))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 12 — APPLICATION DEPLOYMENT FLOW
# ══════════════════════════════════════════════════════════════
add_heading(doc, "12. Application Deployment Flow", level=1)

add_para(doc, "The complete deployment sequence from a developer git push to pods serving traffic:")

add_two_col_table(doc, [
    ("1. Code Push",             "Developer pushes to main branch on GitHub"),
    ("2. Webhook Trigger",        "GitHub fires POST to https://jenkins.ankit.services/github-webhook/"),
    ("3. Jenkins Job Start",      "Jenkins creates an ephemeral agent pod in the jenkins namespace"),
    ("4. Checkout",               "Pipeline clones the repository"),
    ("5. Parallel Build",         "If app/frontend/** changed: build frontend Docker image (v<BUILD_NUMBER>)"),
    ("",                          "If app/backend/** changed: build backend Docker image (v<BUILD_NUMBER>)"),
    ("6. ECR Login & Push",       "aws ecr get-login-password | docker login → docker push versioned + :latest"),
    ("7. kubeconfig Update",      "aws eks update-kubeconfig syncs credentials for todo-tf-cluster-dev"),
    ("8. Rolling Deploy",         "kubectl set image deployment/<name> <container>=<image>:<tag> -n todo"),
    ("9. Rollout Watch",          "kubectl rollout status --timeout=180s — monitors new pod readiness probes"),
    ("10. Auto-rollback (if fail)","kubectl rollout undo reverts to previous ReplicaSet if rollout times out"),
    ("11. Cleanup",               "Docker images removed from agent disk; tfplan file deleted"),
    ("12. Agent Pod Deleted",     "Kubernetes terminates the agent pod after pipeline completion"),
], header=["Step", "Action"])

add_para(doc, "")
add_para(doc, (
    "For infrastructure changes (when infra/ is modified), steps 5-10 are replaced by "
    "Terraform Init → Validate → Plan → Apply. Both flows can happen in the same pipeline "
    "run if a commit touches both app/ and infra/."
))

add_subheading(doc, "12.1  Manual Deployment (deploy.sh)", level=2)
add_para(doc, (
    "For emergency out-of-band deployments, app/ci/deploy.sh provides a seven-step "
    "manual deployment with the same safety guarantees as the pipeline: preflight checks, "
    "ECR push, manifest application in dependency order, and automatic rollback on failure."
))
add_code_block(doc, """\
# Full build + deploy:
cd app && ./ci/deploy.sh

# Manifests only (skip Docker build/push — use existing image):
./ci/deploy.sh --skip-build

# Deploy + also push :latest tag:
./ci/deploy.sh --latest-tag""")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 13 — REQUEST LIFECYCLE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "13. End-to-End Request Lifecycle", level=1)

add_placeholder(doc,
    "Insert Full Request Traffic Flow Diagram Here",
    "Shows user browser → Route53 → ALB (public subnet) → path-based routing → "
    "frontend pod or backend pod (private subnet) → backend → Redis check → RDS. "
    "Reference: docs/diagrams/Full Request Traffic Flow-2026-05-19-085523.png"
)

add_para(doc, "The following trace follows a GET /api/todos request through the full stack:")

add_two_col_table(doc, [
    ("1. DNS Resolution",   "Browser resolves www.ankit.services → Route53 returns ALB DNS name → browser connects to ALB IP"),
    ("2. TLS Handshake",    "ALB terminates TLS 1.3 using ACM cert; HTTP/1.1 or HTTP/2 established with backend"),
    ("3. HTTP→HTTPS",        "Any HTTP:80 request receives 301 Redirect to HTTPS:443 before reaching the application"),
    ("4. ALB Path Routing",  "/api/todos matches the /api/* listener rule → request forwarded to backend target group"),
    ("5. Target Selection",  "ALB selects a healthy backend pod IP (target-type: ip — direct pod routing, no kube-proxy hop)"),
    ("6. Network Policy",    "ALB ENI (in 10.0.0.0/8) is allowed by backend NetworkPolicy; connection accepted"),
    ("7. Backend Receives",  "Express handles GET /api/todos → controller.getTodos()"),
    ("8. Cache Check",       "redis.get('todos:all') — if cache hit, return JSON immediately (sub-ms response)"),
    ("9. DB Query (miss)",   "If cache miss: db.execute('SELECT * FROM todos') → MySQL RDS in private subnet"),
    ("10. Cache Populate",   "redis.set('todos:all', JSON, 'EX', 60) — TTL 60s to avoid stale reads"),
    ("11. Response",         "JSON array of todos returned to ALB → ALB forwards to client over TLS"),
], header=["Step", "Detail"])

add_para(doc, "")
add_subheading(doc, "13.1  Frontend Static Asset Request", level=2)
add_para(doc, (
    "For GET / (or any non-/api path), the ALB default rule forwards to the frontend target group. "
    "The frontend pod runs serve@14 which serves the static files directly. The browser receives "
    "index.html, then fetches app.js and style.css. Once the page loads, app.js makes XHR calls "
    "to /api/todos which follow the backend path above."
))

add_placeholder(doc,
    "Insert ALB Path Routing Diagram Here",
    "Shows ALB HTTPS listener with two rules: Priority 10 — /api/* and /health → backend TG; "
    "Default — / → frontend TG. "
    "Reference: docs/diagrams/ALB Path Routing-2026-05-19-085341.png"
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 14 — SECURITY IMPLEMENTATION
# ══════════════════════════════════════════════════════════════
add_heading(doc, "14. Security Implementation", level=1)

add_subheading(doc, "14.1  IAM Least Privilege", level=2)
add_para(doc, (
    "Every component has the minimum IAM permissions needed and nothing more. The backend pod "
    "can only read its specific Secrets Manager secret (resource-level ARN restriction). "
    "The ALB controller has a detailed policy (alb-controller-policy.json) listing only the "
    "ELB/EC2 API calls it needs. The Cluster Autoscaler role allows only ASG describe/scale "
    "operations. Node roles use AWS-managed policies appropriate to their function."
))
add_para(doc, (
    "Jenkins is the exception: it holds PowerUserAccess + IAMFullAccess to run Terraform. "
    "This is deliberately permissive because Terraform creates IAM roles during infrastructure "
    "provisioning. In a stricter environment, this would be split into a dedicated Terraform "
    "pipeline role with more targeted permissions, or constrained with permission boundaries."
))

add_subheading(doc, "14.2  Secrets Handling", level=2)
add_para(doc, (
    "No sensitive values appear in git at any point. The secret chain is: "
    "Terraform writes credentials to Secrets Manager (read from TF_VAR_rds_password env var) → "
    "Secrets Store CSI Driver reads from Secrets Manager at pod startup (using IRSA) → "
    "Kubernetes Secret synced → env vars in container. "
    "The Kubernetes Secret (todo-secret) is not stored in git; its existence depends entirely "
    "on the CSI driver mounting the volume and syncing the secret."
))

add_subheading(doc, "14.3  Network Isolation", level=2)
bullets = [
    "RDS and ElastiCache: accessible only from EKS node security group. No public endpoint.",
    "EKS nodes: in private subnets, no public IP. Internet access via NAT GW.",
    "ALB: in public subnets, accepts only 443 and 80. All other ports blocked by ALB SG.",
    "Kubernetes NetworkPolicies: backend accepts only ALB VPC CIDR + frontend pods. Frontend accepts only ALB CIDR. Default deny applies to all selected pods.",
    "IMDSv2: http_tokens = required on all EKS node launch template instances — prevents SSRF credential theft from pod-reachable metadata endpoint.",
]
for b in bullets:
    add_bullet(doc, b)

add_subheading(doc, "14.4  TLS/SSL", level=2)
add_para(doc, (
    "TLS is enforced end-to-end at the ALB with policy ELBSecurityPolicy-TLS13-1-2-2021-06. "
    "Traffic between ALB and pods is HTTP over the VPC private network — this is acceptable "
    "because VPC traffic doesn't traverse the public internet and the ALB handles external TLS. "
    "For compliance environments requiring end-to-end TLS, backend containers would need to serve "
    "HTTPS with self-signed or private CA certificates."
))

add_subheading(doc, "14.5  Container Security Context", level=2)
add_code_block(doc, """\
securityContext:
  runAsNonRoot: true          # pod level
  runAsUser: 1000
  fsGroup: 1000

containers:
  securityContext:
    allowPrivilegeEscalation: false
    readOnlyRootFilesystem: false   # serve + dotenv need fs access
    capabilities:
      drop: ["ALL"]""")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 15 — MONITORING AND LOGGING
# ══════════════════════════════════════════════════════════════
add_heading(doc, "15. Monitoring and Logging", level=1)

add_subheading(doc, "15.1  CloudWatch Log Groups", level=2)
add_two_col_table(doc, [
    ("/aws/eks/todo-tf-cluster-dev/cluster",  "EKS control plane logs: api, audit, authenticator, controllerManager, scheduler — 14-day retention"),
    ("/todo/dev/application",                  "General application logs — 14-day retention"),
    ("/todo/dev/backend",                      "Backend-specific logs — 14-day retention"),
    ("/todo/dev/frontend",                     "Frontend-specific logs — 14-day retention"),
], header=["Log Group", "Content & Retention"])

add_subheading(doc, "15.2  RDS Logs", level=2)
add_para(doc, (
    "The RDS instance exports error and slowquery logs to CloudWatch. These are invaluable for "
    "diagnosing N+1 query problems and unexpected error conditions. The slow query log threshold "
    "defaults to MySQL's built-in (10 seconds) — lowering it to 1-2 seconds in production is "
    "recommended."
))

add_subheading(doc, "15.3  Node-level Metrics", level=2)
add_para(doc, (
    "The node IAM role includes CloudWatchAgentServerPolicy, enabling the CloudWatch agent "
    "to push node-level metrics (CPU, memory, disk, network) without additional credential "
    "configuration. HPA relies on the Kubernetes Metrics Server (installed as a CoreDNS dependency) "
    "for pod-level CPU/memory metrics."
))

add_subheading(doc, "15.4  Application-Level Observability", level=2)
add_para(doc, (
    "The backend logs cache hits (✅ Cache hit), cache misses (⏳ Cache miss — querying RDS), "
    "database connection status, Redis connection status, and unhandled errors. "
    "These structured log lines make it straightforward to correlate CloudWatch Logs Insights "
    "queries against observed latency spikes."
))
add_para(doc, (
    "The /health endpoint returns {\"status\": \"OK\", \"version\": \"v4\"} and is used by both "
    "the ALB health check (unhealthy-threshold: 3 at 15-second intervals) and the Kubernetes "
    "liveness/readiness probes."
))

add_placeholder(doc,
    "Insert CloudWatch Monitoring Dashboard Screenshot Here",
    "Recommended dashboard showing: EKS node CPU/memory, backend pod HPA replica count, "
    "ALB request count and 5xx error rate, RDS connections and slow query count, "
    "Redis cache hit ratio (derived from backend logs)."
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 16 — SCALABILITY CONSIDERATIONS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "16. Scalability Considerations", level=1)

add_subheading(doc, "16.1  Pod-level Scaling (HPA)", level=2)
add_para(doc, (
    "The HPA v2 configuration uses dual metrics (CPU and memory) with independent thresholds. "
    "The backend scales between 2 and 10 replicas; the frontend between 2 and 6. "
    "The choice to scale frontend less aggressively reflects its stateless, IO-light nature — "
    "a single serve process can handle many concurrent static file requests."
))

add_subheading(doc, "16.2  Node-level Scaling (Cluster Autoscaler)", level=2)
add_para(doc, (
    "The Cluster Autoscaler watches for pending pods (pods that can't be scheduled due to "
    "insufficient resources) and scales the EKS node group from its minimum (1 in dev, 2 in prod) "
    "up to its maximum (3 in dev). The complementary ASG target-tracking policy also scales "
    "out at 60% average CPU to add capacity proactively."
))
add_para(doc, (
    "A known limitation: Cluster Autoscaler scale-down is conservative — it waits 10 minutes "
    "by default before removing underutilized nodes. PodDisruptionBudgets ensure minAvailable: 1 "
    "is respected during these scale-down node drains."
))

add_subheading(doc, "16.3  Database Scaling", level=2)
add_para(doc, (
    "RDS is configured with max_allocated_storage = allocated_storage * 3 (up to 60 GiB) for "
    "autoscaling storage expansion. Connection pool is limited to 10 per backend pod. "
    "With 10 backend replicas (HPA max), that's 100 concurrent DB connections — within "
    "db.t3.micro's limits but approaching them. For sustained high load, consider RDS Proxy "
    "to multiplex connections, or upgrade the instance class."
))
add_para(doc, (
    "The Redis cache-aside pattern (60-second TTL on todos:all) effectively removes the DB "
    "from the GET /api/todos read path for the vast majority of requests. Invalidation on "
    "write operations (create, update, delete) ensures consistency within one TTL window."
))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 17 — COST OPTIMIZATION
# ══════════════════════════════════════════════════════════════
add_heading(doc, "17. Cost Optimization", level=1)

add_two_col_table(doc, [
    ("Instance Types",          "t3.medium (EKS nodes), db.t3.micro (RDS), cache.t3.micro (Redis) — T-series burstable instances for dev/non-prod workloads"),
    ("Single NAT Gateway",      "One NAT GW in the first AZ vs. one per AZ — saves ~$32/month. Acceptable tradeoff for dev; not recommended for prod."),
    ("ECR Lifecycle Policy",    "Retains only 5 tagged images per repo — prevents unbounded ECR storage cost as builds accumulate"),
    ("CloudWatch Retention",    "14-day log retention in dev vs. 30-day in prod — reduces CloudWatch storage cost proportionally"),
    ("Cluster Autoscaler",      "Scales EKS node group down to 1 node in dev during low-traffic periods — largest cost saving in non-prod environments"),
    ("On-Demand vs Spot",       "Currently ON_DEMAND capacity type. Switching non-backend nodes to Spot instances could reduce node cost by 60-70% for dev workloads."),
    ("RDS Multi-AZ",            "Disabled in dev (rds_multi_az = false) — saves cost of standby replica. Variable-driven for easy prod promotion."),
    ("EBS gp3",                 "gp3 volumes are ~20% cheaper than gp2 at the same IOPS. Used for both EKS node root volumes and RDS storage."),
], header=["Optimization", "Detail"])

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 18 — PRODUCTION BEST PRACTICES IMPLEMENTED
# ══════════════════════════════════════════════════════════════
add_heading(doc, "18. Production Best Practices Implemented", level=1)

practices = [
    ("Zero-Credential IRSA",          "No AWS_ACCESS_KEY_ID or static credentials anywhere. All AWS access via IRSA token exchange."),
    ("Secrets Never in Git",          "Secrets Manager + CSI Driver pattern. No .env files committed. No secret in K8s manifests (base64 is not encryption)."),
    ("Immutable Container Images",    "Every build produces a versioned image tag. Tags are never overwritten (only :latest is mutable for quick rollback references)."),
    ("Rolling Deployments",           "maxUnavailable: 0 ensures at least 2 replicas are always serving traffic during a deploy."),
    ("Automated Rollback",            "Pipeline and deploy.sh both execute kubectl rollout undo on rollout timeout."),
    ("PodDisruptionBudgets",          "Prevents maintenance operations from removing all application replicas simultaneously."),
    ("Pod Anti-Affinity",             "Pods spread across AZs — a single AZ failure doesn't take the application offline."),
    ("ReadOnly Root FS (partial)",    "Enabled where possible; disabled only where process needs runtime file writes (noted in code)."),
    ("IMDSv2 Enforcement",            "http_tokens = required on all nodes prevents SSRF attacks from pods."),
    ("Storage Encryption",            "RDS encrypted at rest. Redis at-rest encryption enabled. EBS node volumes encrypted."),
    ("Log Retention Policies",        "CloudWatch log groups have explicit retention periods to prevent unbounded log accumulation."),
    ("Pre-created Log Groups",        "CloudWatch log groups created before EKS cluster — prevents EKS from failing to log due to missing group."),
    ("Resource Requests & Limits",    "Every container has both set — prevents resource contention and enables HPA to function correctly."),
    ("Health Probes",                 "Liveness and readiness probes on all containers — Kubernetes knows when a pod is actually healthy, not just started."),
    ("Terraform State Locking",       "Native S3 lock file prevents concurrent Terraform runs from corrupting state."),
    ("CORS Configuration",            "Backend ALLOWED_ORIGIN restricted to https://www.ankit.services — no wildcard in production."),
    ("changeSet-Gated Pipeline",      "Only changed components are built/deployed — reduces blast radius and pipeline duration."),
    ("DB Connection Retry",           "Backend retries DB connection 15 times with 5-second backoff — handles the RDS cold-start latency on first pod launch."),
]
add_two_col_table(doc, practices, header=["Practice", "Implementation"])

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 19 — CHALLENGES FACED AND SOLUTIONS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "19. Challenges Faced and Solutions", level=1)

challenges = [
    (
        "EKS Cluster Logging Race Condition",
        "EKS control plane tried to write logs during cluster creation, but the CloudWatch log group "
        "didn't exist yet — causing a transient error in cluster creation.",
        "Created the CloudWatch log group as a separate Terraform resource (cloudwatch module) with an "
        "explicit depends_on in the EKS module. The log group now pre-exists before the cluster resource is created."
    ),
    (
        "ALB Controller Security Group Rule",
        "After switching to a custom launch template (to enforce IMDSv2 and custom node SG), EKS no longer "
        "auto-attached the cluster security group to nodes. Control plane couldn't reach kubelets on port 10250 "
        "— kubectl exec and logs commands failed.",
        "Added three explicit aws_security_group_rule resources in the EKS module to restore control plane → node "
        "communication: kubelet (10250), admission webhooks (443), and ephemeral ports (1025-65535)."
    ),
    (
        "Secrets Store CSI Driver: SecretProviderClass CRD Not Found",
        "The SecretProviderClass resource applied before the Secrets Store CSI Driver Helm release was fully "
        "installed — CRD didn't exist yet, giving a 'no matches for kind SecretProviderClass' error.",
        "Both Helm releases (CSI driver + AWS provider) are now managed by Terraform with depends_on = [module.eks], "
        "ensuring they are installed before any kubectl apply of application manifests."
    ),
    (
        "ALB TargetGroupBinding vs Ingress Conflict",
        "Initially used only the Ingress resource to manage the ALB. Later switched to pre-creating ALB/TGs in "
        "Terraform and using TargetGroupBinding CRDs to bind them. Both approaches tried to manage the same ALB "
        "at the same time — annotations conflicted and the controller threw ownership errors.",
        "Settled on a hybrid: Terraform owns the ALB and target groups; Ingress annotation references the existing "
        "ALB ARN (alb.ingress.kubernetes.io/load-balancer-arn) to attach to it instead of creating a new one. "
        "TargetGroupBinding CRDs handle target registration."
    ),
    (
        "Jenkins Agent Can't Connect to Jenkins Controller",
        "Agent pods were connecting to the external Jenkins URL (jenkins.ankit.services) which routes through "
        "the ALB → back into the cluster. This hairpin causes connection timeouts in some VPC configurations.",
        "Set jenkinsUrl to the internal cluster DNS name (http://jenkins.jenkins.svc.cluster.local:8080) in "
        "jenkins-values.yaml. Agent pods resolve this via CoreDNS and connect directly within the cluster."
    ),
    (
        "Terraform IRSA: oidc_issuer_host Ordering Problem",
        "The IRSA trust policy in secrets-manager/main.tf needed the OIDC issuer host string, which is only "
        "available after the OIDC provider is created in the EKS module. Terraform couldn't determine the value "
        "at plan time, causing a 'known after apply' plan failure in modules downstream of EKS.",
        "Passed oidc_issuer_host as an explicit output from the EKS module and threaded it as an input variable "
        "to the secrets_manager and autoscalling modules. Terraform resolves it at apply time after EKS is done."
    ),
    (
        "Backend Pod Startup: DB Not Ready",
        "On first deploy to a cold cluster, the RDS instance takes 60-90 seconds to become available after "
        "the Terraform apply. Backend pods started immediately and crashed-loop waiting for MySQL.",
        "The db.js module implements a retry loop: 15 attempts with 5-second delays. Kubernetes readiness probes "
        "keep the pod out of the load balancer rotation until the DB connects successfully and the app starts "
        "listening on port 3000."
    ),
    (
        "Redis transit_encryption_enabled Conflict",
        "Setting transit_encryption_enabled = true on ElastiCache requires the client to connect with TLS. "
        "The ioredis client was connecting without TLS, causing all Redis operations to fail with 'NOAUTH' "
        "or SSL_ERROR_RX_UNEXPECTED_HELLO errors.",
        "Disabled transit encryption (set to false) for now — acceptable within the VPC's private subnet. "
        "The tradeoff is noted in the Redis module with a comment. Enabling it requires adding tls: {} "
        "to the ioredis client configuration."
    ),
]

for challenge, problem, solution in challenges:
    add_subheading(doc, challenge, level=2)
    p = doc.add_paragraph()
    r = p.add_run("Problem: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)
    p.add_run(problem)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Solution: ")
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x00, 0x72, 0x00)
    p2.add_run(solution)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 20 — KEY LEARNINGS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "20. Key Learnings", level=1)

add_subheading(doc, "20.1  Technical Learnings", level=2)
tech_learnings = [
    "IRSA is the correct way to grant AWS permissions to Kubernetes workloads. Static access keys in env vars or secrets are an anti-pattern even in development.",
    "Terraform's depends_on doesn't work for provider configuration — you must use data sources (aws_eks_cluster + aws_eks_cluster_auth) to initialize the Helm provider after the cluster exists.",
    "EKS managed node groups and custom launch templates interact in non-obvious ways. Setting vpc_security_group_ids in a launch template disables the cluster's default security group auto-attachment — you must manually re-add the control plane communication rules.",
    "The Secrets Store CSI Driver requires a pod volume mount to trigger secret syncing. A secret is only populated (and the Kubernetes Secret created) once a pod that references the SecretProviderClass is scheduled. This means the todo-secret Kubernetes Secret doesn't exist until the first backend pod starts.",
    "Redis cache invalidation on every write (del key) is simpler and more correct than cache update. For a read-heavy, write-light workload like this, it introduces negligible latency.",
    "target-type: ip (direct pod routing) in the ALB requires that the ALB security group can reach pod IP addresses — which are VPC-native IPs with VPC CNI. The NetworkPolicy ipBlock must include the ALB ENI subnets (10.0.0.0/8) not just the cluster CIDR.",
]
for l in tech_learnings:
    add_bullet(doc, l)

add_subheading(doc, "20.2  Operational Learnings", level=2)
ops_learnings = [
    "Rolling update with maxUnavailable: 0 is the safest default for production services. The brief period of running more pods (maxSurge: 1) is acceptable; having zero pods serving requests is not.",
    "kubectl rollout status as a pipeline gate is essential. Without it, the pipeline would report success even if the new pods immediately crash-looped after deploy.",
    "Pre-creating CloudWatch log groups is a small but important operational detail. EKS control plane logs failing silently during a cluster incident (because the log group didn't exist yet) is a bad time to discover this.",
    "Changeset-based pipeline gating dramatically improves developer experience. Pushing a CSS fix should not trigger an 8-minute infrastructure pipeline.",
    "The deploy.sh script is worth maintaining as a fallback. In a real incident where the Jenkins pipeline itself is broken, having a tested manual deploy path prevents a single point of failure in the deployment toolchain.",
]
for l in ops_learnings:
    add_bullet(doc, l)

add_subheading(doc, "20.3  Architectural Learnings", level=2)
arch_learnings = [
    "Placing all Terraform modules under a flat root module (vs. separate state files per module) simplifies cross-module references (like passing oidc_provider_arn between modules) but creates a large blast radius if apply fails mid-way. For very large projects, split state boundaries around lifecycle: foundational (VPC/EKS) vs. application (RDS/ECR/secrets).",
    "The TargetGroupBinding CRD approach (Terraform creates ALB/TGs, K8s binds services to them) is more robust than letting the ALB Controller manage everything. It separates infrastructure lifecycle (Terraform) from service routing (Kubernetes) cleanly.",
    "Single NAT Gateway is a real cost optimization that carries operational risk. For any service with an SLA, a per-AZ NAT Gateway is table stakes.",
    "Using the same ACM certificate for both the app and Jenkins simplifies cert management at the cost of a slightly wider trust domain. A wildcard cert (*.ankit.services) would have been cleaner.",
]
for l in arch_learnings:
    add_bullet(doc, l)

add_subheading(doc, "20.4  DevOps Learnings", level=2)
devops_learnings = [
    "DevOps is primarily a set of practices, not tools. The pipeline, IRSA, and secret management patterns here matter more than which specific tool implements them.",
    "Every environment variable that a pod needs at runtime should be traceable to either a ConfigMap (non-sensitive config) or a Secret populated from Secrets Manager (sensitive config). Nothing else.",
    "Infrastructure as Code only delivers its full value when the state is the source of truth. Making manual AWS console changes and not reflecting them in Terraform creates drift — and drift causes the next terraform apply to revert or fail on the manual changes.",
    "Testing rollback paths proactively (by intentionally deploying a broken image) reveals assumptions in the health probe and timeout configuration that aren't apparent from reading the manifests.",
]
for l in devops_learnings:
    add_bullet(doc, l)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 21 — FUTURE IMPROVEMENTS
# ══════════════════════════════════════════════════════════════
add_heading(doc, "21. Future Improvements", level=1)

improvements = [
    ("GitOps with ArgoCD",          "Replace the kubectl set image pipeline with ArgoCD + image updater. Git becomes the source of truth for Kubernetes state; ArgoCD continuously reconciles cluster state to git."),
    ("Per-AZ NAT Gateways",         "Add a NAT GW per availability zone for HA and to eliminate cross-AZ NAT data transfer costs in production."),
    ("RDS Multi-AZ + RDS Proxy",    "Enable Multi-AZ for RDS in production. Add RDS Proxy to pool connections, especially when backend HPA scales to 10 replicas."),
    ("Prometheus + Grafana",         "Replace CloudWatch with the Prometheus / Grafana stack (kube-prometheus-stack Helm chart) for richer metrics, alerting, and dashboard customization."),
    ("Distributed Tracing",          "Add OpenTelemetry instrumentation to the backend for request-level tracing with Jaeger or AWS X-Ray."),
    ("End-to-End TLS",               "Enable transit_encryption_enabled on ElastiCache and update ioredis config. Consider a service mesh (Istio / Linkerd) for mTLS between all pods."),
    ("Separate Terraform State Files","Split into foundation (VPC/EKS) and application (RDS/Redis/ECR/secrets) statefiles to reduce blast radius and allow parallel development."),
    ("Velero for Backup",            "Add Velero for Kubernetes resource and persistent volume backup. EBS snapshots + RDS automated backups cover data, but Velero covers the Kubernetes object state."),
    ("Spot Instances for Workers",   "Use a mixed node group (on-demand + Spot) for cost savings. PodDisruptionBudgets are already in place to handle Spot interruption drains gracefully."),
    ("Container Image Scanning",     "Enable ECR enhanced scanning (Inspector) on the ECR repositories to catch CVEs in base images before they reach production."),
    ("RBAC Hardening",               "Currently relying on IRSA at the pod level without explicit Kubernetes RBAC for developer users. Add RBAC ClusterRoles and aws-auth ConfigMap entries for least-privilege kubectl access."),
    ("Pipeline Test Stage",          "Add a stage before Deploy that runs backend integration tests against a test database — currently there are no automated tests in the pipeline."),
]
add_two_col_table(doc, improvements, header=["Improvement", "Rationale"])

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 22 — DEPLOYMENT COMMANDS REFERENCE
# ══════════════════════════════════════════════════════════════
add_heading(doc, "22. Deployment Commands Reference", level=1)

add_subheading(doc, "22.1  Prerequisites", level=2)
add_code_block(doc, """\
# Required tools:
aws --version          # AWS CLI v2
terraform --version    # 1.9+
kubectl version        # 1.30
helm version           # 3.x
docker --version       # 24+

# AWS credentials configured:
aws sts get-caller-identity""")

add_subheading(doc, "22.2  Terraform Bootstrap (one-time)", level=2)
add_code_block(doc, """\
# Create S3 state bucket (one-time):
aws s3 mb s3://todo-tf-state-668076964228 --region us-east-1
aws s3api put-bucket-versioning \\
    --bucket todo-tf-state-668076964228 \\
    --versioning-configuration Status=Enabled

# Initialize with dev backend:
cd infra/terraform
terraform init -backend-config=environments/dev/backend.tf

# Set sensitive variable:
export TF_VAR_rds_password='your-secure-password'

# Plan and apply:
terraform plan -var-file=environments/dev/terraform.tfvars -out=tfplan
terraform apply tfplan""")

add_subheading(doc, "22.3  Connect kubectl to EKS", level=2)
add_code_block(doc, """\
aws eks update-kubeconfig \\
    --region us-east-1 \\
    --name todo-tf-cluster-dev

kubectl get nodes
kubectl get pods -A""")

add_subheading(doc, "22.4  Install Jenkins (Helm)", level=2)
add_code_block(doc, """\
# Create namespace:
kubectl create namespace jenkins

# Build and push Jenkins agent image:
docker build -f app/docker/Dockerfile.jenkins-agent \\
    -t 668076964228.dkr.ecr.us-east-1.amazonaws.com/jenkins-agent:latest .
aws ecr get-login-password --region us-east-1 | \\
    docker login --username AWS --password-stdin 668076964228.dkr.ecr.us-east-1.amazonaws.com
docker push 668076964228.dkr.ecr.us-east-1.amazonaws.com/jenkins-agent:latest

# Install Jenkins:
helm repo add jenkins https://charts.jenkins.io
helm repo update
helm install jenkins jenkins/jenkins \\
    --namespace jenkins \\
    --values infra/ci-cd/jenkins-values.yaml \\
    --set controller.admin.password=<your-admin-password>""")

add_subheading(doc, "22.5  Apply Kubernetes Manifests", level=2)
add_code_block(doc, """\
# Order matters — namespace first, then config, then workloads:
kubectl apply -f app/k8s/namespace.yaml
kubectl apply -f app/k8s/serviceaccount.yaml  -n todo
kubectl apply -f app/k8s/secretproviderclass.yaml -n todo
kubectl apply -f app/k8s/configmap.yaml        -n todo
kubectl apply -f app/k8s/deployment.yaml       -n todo
kubectl apply -f app/k8s/service.yaml          -n todo
kubectl apply -f app/k8s/ingress.yaml          -n todo
kubectl apply -f app/k8s/hpa.yaml              -n todo
kubectl apply -f app/k8s/pdb.yaml              -n todo
kubectl apply -f app/k8s/networkpolicy.yaml    -n todo
kubectl apply -f app/k8s/targetgroupbinding.yaml -n todo""")

add_subheading(doc, "22.6  Manual Image Deploy", level=2)
add_code_block(doc, """\
# Build and push backend manually:
export IMAGE_TAG=v$(date +%Y%m%d-%H%M%S)
docker build -f app/docker/Dockerfile.backend \\
    -t 668076964228.dkr.ecr.us-east-1.amazonaws.com/todo-backend:${IMAGE_TAG} \\
    app/backend/
aws ecr get-login-password --region us-east-1 | \\
    docker login --username AWS --password-stdin 668076964228.dkr.ecr.us-east-1.amazonaws.com
docker push 668076964228.dkr.ecr.us-east-1.amazonaws.com/todo-backend:${IMAGE_TAG}

# Rolling deploy:
kubectl set image deployment/todo-backend \\
    todo-backend=668076964228.dkr.ecr.us-east-1.amazonaws.com/todo-backend:${IMAGE_TAG} \\
    -n todo
kubectl rollout status deployment/todo-backend -n todo --timeout=180s""")

add_subheading(doc, "22.7  Useful Debugging Commands", level=2)
add_code_block(doc, """\
# Pod status:
kubectl get pods -n todo -o wide
kubectl describe pod <pod-name> -n todo
kubectl logs <pod-name> -n todo --previous

# HPA status:
kubectl get hpa -n todo

# Ingress status:
kubectl get ingress -n todo
kubectl describe ingress todo-ingress -n todo

# Secrets Manager:
aws secretsmanager get-secret-value \\
    --secret-id todo-dev-app-credentials \\
    --region us-east-1

# Terraform outputs:
terraform -chdir=infra/terraform output

# Force rollback:
kubectl rollout undo deployment/todo-backend -n todo
kubectl rollout undo deployment/todo-frontend -n todo""")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 23 — TROUBLESHOOTING CHEAT SHEET
# ══════════════════════════════════════════════════════════════
add_heading(doc, "23. Troubleshooting Cheat Sheet", level=1)

add_two_col_table(doc, [
    ("Pods stuck in Pending",
     "Check: kubectl describe pod <name> -n todo → Events section.\n"
     "Cause: Insufficient node capacity (scale-up in progress), unschedulable due to anti-affinity, PVC not bound.\n"
     "Fix: Wait for Cluster Autoscaler or increase node max; check PVC status."),

    ("Pods stuck in Init:CrashLoopBackOff",
     "Check: kubectl logs <pod> -n todo --previous\n"
     "Cause: Secrets Store CSI volume mount failing (IRSA misconfiguration, secret name mismatch).\n"
     "Fix: Verify service account annotation, IAM role trust policy, SecretProviderClass object names."),

    ("Backend returns 500 on /api/todos",
     "Check: Backend logs for DB connection errors.\n"
     "Cause: DB_HOST wrong, RDS security group blocking, DB not yet accepting connections.\n"
     "Fix: kubectl exec into backend pod → ping $DB_HOST. Verify SG rule allows EKS node SG on port 3306."),

    ("ALB health check failing",
     "Check: Target group target health in AWS console.\n"
     "Cause: /health endpoint not responding (pod crash), port mismatch, security group blocking ALB→node.\n"
     "Fix: Confirm ALB SG → EKS node SG allows port 3000. Check pod readiness probe."),

    ("kubectl commands fail after eks update-kubeconfig",
     "Check: aws sts get-caller-identity — confirm correct IAM identity.\n"
     "Cause: aws-auth ConfigMap doesn't include IAM user/role, or IRSA role not applied.\n"
     "Fix: kubectl edit configmap aws-auth -n kube-system to add IAM principal."),

    ("Terraform apply fails: 'AccessDenied'",
     "Check: Error message for specific API call and resource ARN.\n"
     "Cause: IAM role/user lacks permission for that specific API.\n"
     "Fix: Add the required policy action. Use CloudTrail to identify the exact denied call."),

    ("Jenkins agent pods not spawning",
     "Check: kubectl get pods -n jenkins. Describe controller pod logs.\n"
     "Cause: Jenkins Kubernetes cloud config incorrect, IRSA not applied, ECR image pull failing.\n"
     "Fix: Verify agent image exists in ECR. Check imagePullPolicy + ECR permissions on node role."),

    ("Redis connection fails in backend",
     "Check: Backend logs for 'Redis error' messages.\n"
     "Cause: REDIS_HOST secret not populated, wrong host, security group blocking, TLS mismatch.\n"
     "Fix: kubectl exec into pod → env | grep REDIS_HOST. Test: redis-cli -h $REDIS_HOST ping."),

    ("HPA stuck at minReplicas despite high load",
     "Check: kubectl describe hpa -n todo — look for 'FailedGetScale' events.\n"
     "Cause: Metrics Server not running, resource requests not set on containers.\n"
     "Fix: kubectl get pods -n kube-system | grep metrics-server. Verify resource requests in deployment spec."),

    ("Terraform: 'state lock' error",
     "Check: Error message includes lock info (operation, who, when).\n"
     "Cause: Previous terraform apply crashed mid-run without releasing the S3 lock file.\n"
     "Fix: If confident the previous run is dead: terraform force-unlock <lock-id>"),
], header=["Symptom", "Cause & Resolution"])

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 24 — CONCLUSION
# ══════════════════════════════════════════════════════════════
add_heading(doc, "24. Conclusion", level=1)

add_para(doc, (
    "This project demonstrates that a complete, production-hardened cloud-native deployment is achievable "
    "with a relatively small, well-structured codebase. The key design decisions — IRSA over static "
    "credentials, Secrets Manager over git-stored secrets, changeset-gated pipelines over full rebuilds "
    "on every commit, zero-downtime rolling deployments with automated rollback — are not complexity for "
    "its own sake. They reflect patterns that make the system safer, more auditable, and easier to "
    "operate at 2am when something goes wrong."
))
add_para(doc, (
    "The infrastructure is entirely recreatable from the Terraform codebase in a clean AWS account. "
    "The application is deployable from git through the Jenkins pipeline without any manual steps. "
    "Every credential is either ephemeral (IRSA) or stored in Secrets Manager. Every component "
    "scales automatically within defined bounds."
))
add_para(doc, (
    "The areas identified for future improvement — GitOps adoption, per-AZ NAT Gateways, RDS Proxy, "
    "Prometheus/Grafana — are genuine production hardening steps, not academic wishlist items. They "
    "represent the delta between 'production-ready for a small team' and 'enterprise-production-ready "
    "for a large team with strict SLAs.' The foundation built here makes those additions straightforward, "
    "not disruptive."
))
add_para(doc, (
    "The application is live and accessible at https://www.ankit.services."
))

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 25 — REFERENCES
# ══════════════════════════════════════════════════════════════
add_heading(doc, "25. References", level=1)

refs = [
    ("Amazon EKS Documentation",              "https://docs.aws.amazon.com/eks/latest/userguide/"),
    ("AWS Load Balancer Controller",           "https://kubernetes-sigs.github.io/aws-load-balancer-controller/"),
    ("Secrets Store CSI Driver",               "https://secrets-store-csi-driver.sigs.k8s.io/"),
    ("AWS Secrets Manager Provider for CSI",   "https://github.com/aws/secrets-store-csi-driver-provider-aws"),
    ("Terraform AWS Provider",                  "https://registry.terraform.io/providers/hashicorp/aws/latest/docs"),
    ("Kubernetes HPA v2",                       "https://kubernetes.io/docs/tasks/run-application/horizontal-pod-autoscale/"),
    ("Cluster Autoscaler for AWS",              "https://github.com/kubernetes/autoscaler/tree/master/cluster-autoscaler/cloudprovider/aws"),
    ("Jenkins Kubernetes Plugin",               "https://plugins.jenkins.io/kubernetes/"),
    ("EKS Best Practices Guide",                "https://aws.github.io/aws-eks-best-practices/"),
    ("Docker Security Best Practices",          "https://docs.docker.com/develop/security-best-practices/"),
    ("AWS VPC Design — re:Invent",              "https://aws.amazon.com/blogs/aws/re-invent-2023-vpc-design/"),
    ("IRSA Deep Dive",                          "https://aws.amazon.com/blogs/opensource/introducing-fine-grained-iam-roles-service-accounts/"),
]
add_two_col_table(doc, refs, header=["Resource", "URL"])

# ══════════════════════════════════════════════════════════════
# SECTION 26 — PROMETHEUS & GRAFANA MONITORING STACK
# ══════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, "26. Prometheus & Grafana Monitoring Stack", level=1)

add_note_box(doc,
    "This stack is fully implemented and live. It was listed as a 'future improvement' in Section 21 "
    "before the implementation was completed. The kube-prometheus-stack Helm release is managed by "
    "Terraform and Grafana is accessible at https://grafana.ankit.services.",
    "STATUS"
)

add_subheading(doc, "26.1  Installation via Terraform", level=2)
add_para(doc, (
    "The kube-prometheus-stack is deployed as a Terraform helm_release resource in infra/terraform/main.tf "
    "(chart version 58.7.2 from prometheus-community). Terraform installs it into the monitoring namespace "
    "(created automatically) after the EKS cluster is ready. This guarantees the monitoring stack exists "
    "before any application workloads are deployed."
))
add_code_block(doc, """\
resource "helm_release" "kube_prometheus_stack" {
  name             = "kube-prometheus-stack"
  namespace        = "monitoring"
  create_namespace = true
  repository       = "https://prometheus-community.github.io/helm-charts"
  chart            = "kube-prometheus-stack"
  version          = "58.7.2"

  disable_openapi_validation = true
  wait                       = false

  # Grafana admin password (supply via TF_VAR_grafana_admin_password)
  set { name = "grafana.adminPassword"; value = var.grafana_admin_password }

  # Expose Grafana via ALB Ingress at grafana.ankit.services
  set { name = "grafana.ingress.enabled";                          value = "true"          }
  set { name = "grafana.ingress.ingressClassName";                 value = "alb"           }
  set { name = "grafana.ingress.hosts[0]";                        value = "grafana.ankit.services" }
  set { name = "grafana.ingress.annotations.alb...scheme";        value = "internet-facing"}
  set { name = "grafana.ingress.annotations.alb...target-type";   value = "ip"            }

  # Dashboard sidecar — auto-imports ConfigMaps with grafana_dashboard="1" label
  set { name = "grafana.sidecar.dashboards.enabled"; value = "true"             }
  set { name = "grafana.sidecar.dashboards.label";   value = "grafana_dashboard"}

  # Discover ServiceMonitors across ALL namespaces (not just monitoring)
  set { name = "prometheus.prometheusSpec.serviceMonitorSelectorNilUsesHelmValues"; value = "false" }
  set { name = "prometheus.prometheusSpec.podMonitorSelectorNilUsesHelmValues";     value = "false" }

  depends_on = [module.eks]
}""")

add_para(doc, (
    "The two boolean flags serviceMonitorSelectorNilUsesHelmValues = false and "
    "podMonitorSelectorNilUsesHelmValues = false are critical. Without them, Prometheus only "
    "discovers ServiceMonitors in its own namespace (monitoring). Setting both to false enables "
    "cluster-wide ServiceMonitor discovery — allowing the todo namespace ServiceMonitor to be scraped."
))

add_subheading(doc, "26.2  What the Stack Deploys", level=2)
add_two_col_table(doc, [
    ("Prometheus Operator",   "Watches ServiceMonitor/PodMonitor CRDs and configures Prometheus dynamically"),
    ("Prometheus",            "Time-series metrics database — scrapes all configured targets every 15s"),
    ("Grafana",               "Dashboard UI — exposed at https://grafana.ankit.services"),
    ("AlertManager",          "Alert routing and notification (Slack, PagerDuty, email — configurable)"),
    ("kube-state-metrics",    "Exports Kubernetes object state as Prometheus metrics (pod counts, status, restarts)"),
    ("node-exporter",         "Host-level metrics per node: CPU, memory, disk I/O, network"),
], header=["Component", "Purpose"])

add_subheading(doc, "26.3  Route53 — Grafana Subdomain Record", level=2)
add_para(doc, (
    "The Route53 module manages four A alias records. In addition to ankit.services (apex), "
    "www.ankit.services (application), and jenkins.ankit.services, it also creates "
    "grafana.ankit.services pointing to the ALB created by the Kubernetes Load Balancer Controller "
    "for the Grafana Ingress. The ALB DNS name and zone ID are supplied as Terraform variables "
    "(grafana_alb_dns_name, grafana_alb_zone_id) and set in environments/dev/terraform.tfvars "
    "after retrieving them from the Kubernetes ingress status."
))
add_code_block(doc, """\
# Retrieve Grafana ALB DNS after stack is deployed:
kubectl get ingress kube-prometheus-stack-grafana -n monitoring \\
  -o jsonpath='{.status.loadBalancer.ingress[0].hostname}'

# Set in environments/dev/terraform.tfvars:
grafana_alb_dns_name = "k8s-monitori-kubeprom-3ccd051ce4-1913664347.us-east-1.elb.amazonaws.com"
grafana_alb_zone_id  = "Z35SXDOTRQ7X7K"   # fixed us-east-1 ELB zone ID

# Then re-apply Terraform to create the Route53 record:
terraform apply -var-file=environments/dev/terraform.tfvars""")

add_subheading(doc, "26.4  ServiceMonitor — app/k8s/monitoring/servicemonitor.yaml", level=2)
add_para(doc, (
    "The ServiceMonitor CRD (provided by the Prometheus Operator) tells Prometheus to scrape "
    "the todo-backend-svc service in the todo namespace at the /metrics path every 15 seconds."
))
add_code_block(doc, """\
apiVersion: monitoring.coreos.com/v1
kind: ServiceMonitor
metadata:
  name: todo-backend
  namespace: monitoring
  labels:
    release: kube-prometheus-stack   # must match Prometheus operator selector label
spec:
  namespaceSelector:
    matchNames: [todo]
  selector:
    matchLabels:
      app: todo-backend              # selects todo-backend-svc
  endpoints:
  - port: http                       # named port in the service spec
    path: /metrics
    interval: 15s""")

add_subheading(doc, "26.5  Prometheus RBAC — app/k8s/monitoring/prometheus-rbac.yaml", level=2)
add_para(doc, (
    "Prometheus runs in the monitoring namespace but needs permission to read service discovery "
    "data (services, endpoints, pods) from the todo namespace. The Role and RoleBinding grant "
    "exactly those permissions to the Prometheus Operator and Prometheus service accounts. "
    "Without this, Prometheus cannot discover backend pod endpoints and ServiceMonitor scraping silently fails."
))
add_code_block(doc, """\
apiVersion: rbac.authorization.k8s.io/v1
kind: Role
metadata:
  name: prometheus-todo-access
  namespace: todo        # grants permission IN the todo namespace
rules:
- apiGroups: [""]
  resources: ["services", "endpoints", "pods"]
  verbs:     ["get", "list", "watch"]
---
kind: RoleBinding
subjects:
- kind: ServiceAccount
  name: kube-prometheus-stack-operator     # prometheus operator
  namespace: monitoring
- kind: ServiceAccount
  name: kube-prometheus-stack-prometheus   # prometheus itself
  namespace: monitoring""")

add_subheading(doc, "26.6  Additional Scrape Config — app/k8s/monitoring/additional-scrape-secret.yaml", level=2)
add_para(doc, (
    "A Kubernetes Secret in the monitoring namespace provides an additional Prometheus scrape config "
    "that uses Kubernetes service-discovery (role: endpoints) to find backend pod IPs dynamically. "
    "This complements the ServiceMonitor and adds rich per-pod labels (pod, namespace, service, container) "
    "to all scraped metrics — enabling pod-level filtering in Grafana dashboards."
))

add_subheading(doc, "26.7  Grafana Dashboard — app/k8s/monitoring/grafana-dashboard-configmap.yaml", level=2)
add_para(doc, (
    "A pre-built Grafana dashboard is delivered as a Kubernetes ConfigMap with the label "
    "grafana_dashboard: \"1\". The Grafana sidecar (enabled via Terraform) watches for ConfigMaps "
    "with this label across all namespaces and automatically imports them into Grafana at startup — "
    "no manual import required."
))
add_two_col_table(doc, [
    ("Dashboard UID",         "todo-app"),
    ("Auto-refresh",          "Every 30 seconds"),
    ("Default time range",    "Last 1 hour"),
    ("Panel count",           "12 panels"),
], header=["Property", "Value"])

add_para(doc, "")
add_two_col_table(doc, [
    ("Request Rate (req/s)",        "sum(rate(http_requests_total[2m])) by (route)             — requests per second per API route"),
    ("Error Rate (5xx %)",          "100 * sum(rate(http_requests_total{status_code=~\"5..\"}[2m])) / sum(...)  — server error percentage"),
    ("P95 Latency (s)",             "histogram_quantile(0.95, ...) by (le, route)              — 95th percentile response time"),
    ("Backend Pods Running",        "count(kube_pod_status_ready{namespace=\"todo\", pod=~\"todo-backend.*\"})"),
    ("Backend Memory Usage (MB)",   "container_memory_working_set_bytes per pod / 1024 / 1024"),
    ("Frontend Pods Running",       "count(kube_pod_status_ready{namespace=\"todo\", pod=~\"todo-frontend.*\"})"),
    ("Frontend Memory Usage (MB)",  "container_memory_working_set_bytes per pod / 1024 / 1024"),
    ("Backend CPU Usage (cores)",   "rate(container_cpu_usage_seconds_total{pod=~\"todo-backend.*\"}[2m])"),
    ("Pod Restarts",                "kube_pod_container_status_restarts_total by pod           — crash indicator"),
    ("Todo Operations by Method",   "sum(rate(http_requests_total{route=~\"/api/todos.*\"}[2m])) by (method)"),
    ("4xx Client Error Rate",       "sum(rate(http_requests_total{status_code=~\"4..\"}[2m])) by (route)"),
    ("Cache Hit Rate",              "rate(cache_hits_total[2m]) / (rate(cache_hits_total[2m]) + rate(cache_misses_total[2m]))"),
], header=["Panel", "PromQL / Description"])

add_subheading(doc, "26.8  Applying Monitoring Manifests", level=2)
add_para(doc, (
    "The Jenkins pipeline applies all four files in the monitoring/ directory at the end of "
    "the Apply K8s Manifests stage (kubectl apply -f app/k8s/monitoring/). For manual apply:"
))
add_code_block(doc, """\
# Apply in dependency order:
kubectl apply -f app/k8s/monitoring/prometheus-rbac.yaml
kubectl apply -f app/k8s/monitoring/servicemonitor.yaml
kubectl apply -f app/k8s/monitoring/additional-scrape-secret.yaml
kubectl apply -f app/k8s/monitoring/grafana-dashboard-configmap.yaml

# Verify Prometheus is scraping the backend:
kubectl port-forward -n monitoring svc/kube-prometheus-stack-prometheus 9090:9090
# Then open http://localhost:9090/targets and look for todo-backend

# Access Grafana:
open https://grafana.ankit.services
# Default login: admin / <value of TF_VAR_grafana_admin_password>""")

add_subheading(doc, "26.9  Local Monitoring (Docker Compose)", level=2)
add_para(doc, (
    "The docker-compose.yaml in app/docker/ also runs Prometheus and Grafana locally for "
    "development-time metrics testing. Prometheus (port 9090) scrapes backend:3000/metrics every 15 "
    "seconds using the static config in app/docker/prometheus.yml. Grafana (port 3001) can be "
    "connected to the local Prometheus datasource to preview the same dashboards before deploying to EKS."
))
add_code_block(doc, """\
# Start full local stack including monitoring:
docker compose -f app/docker/docker-compose.yaml up

# Prometheus UI:    http://localhost:9090
# Grafana UI:       http://localhost:3001  (admin / admin)

# In Grafana: Settings → Data Sources → Add → Prometheus → URL: http://prometheus:9090
# Import dashboard from: app/k8s/monitoring/grafana-dashboard-configmap.yaml (copy todo-app.json)""")

add_subheading(doc, "26.10  Troubleshooting the Monitoring Stack", level=2)
add_two_col_table(doc, [
    ("ServiceMonitor not found by Prometheus",
     "Check: serviceMonitorSelectorNilUsesHelmValues = false in Terraform. "
     "Verify label release: kube-prometheus-stack on the ServiceMonitor. "
     "Fix: re-apply Terraform with the correct helm set values."),
    ("/metrics endpoint returns 404",
     "The backend pod is running an older image built before metrics.js was added. "
     "Fix: trigger a backend rebuild in Jenkins or docker push a new image."),
    ("Grafana dashboard shows 'No data'",
     "Prometheus is not scraping the backend. Check Prometheus targets at "
     "http://localhost:9090/targets (port-forward). Verify prometheus-rbac.yaml is applied in the todo namespace."),
    ("grafana.ankit.services unreachable",
     "The Grafana ALB DNS name / zone ID in terraform.tfvars may not yet be set. "
     "Fix: kubectl get ingress -n monitoring, get the hostname, add to tfvars, re-apply Terraform."),
    ("Cache Hit Rate panel shows NaN",
     "No write operations have occurred yet (no cache_misses_total increments). "
     "Create a todo item to trigger a cache miss, then GET /api/todos to get a cache hit."),
], header=["Symptom", "Cause & Fix"])

# ══════════════════════════════════════════════════════════════
# SECTION 27 — BACKEND PROMETHEUS METRICS INSTRUMENTATION
# ══════════════════════════════════════════════════════════════
add_page_break(doc)
add_heading(doc, "27. Backend Prometheus Metrics Instrumentation", level=1)

add_para(doc, (
    "The backend is fully instrumented with Prometheus metrics using the prom-client npm library. "
    "This instrumentation is the data source for all Grafana dashboard panels. The metrics module "
    "(app/backend/src/metrics.js) and the middleware in app/backend/src/app.js work together "
    "to capture four categories of signal: HTTP request performance, cache effectiveness, "
    "application errors, and Node.js process health."
))

add_subheading(doc, "27.1  Metrics Module — app/backend/src/metrics.js", level=2)
add_code_block(doc, """\
const client = require('prom-client');

// Default metrics: Node.js process CPU, heap, GC, event-loop lag
client.collectDefaultMetrics();

// HTTP request latency histogram — enables percentile queries (p50, p95, p99)
const httpRequestDuration = new client.Histogram({
  name:       'http_request_duration_seconds',
  help:       'Duration of HTTP requests in seconds',
  labelNames: ['method', 'route', 'status_code'],
  buckets:    [0.005, 0.01, 0.025, 0.05, 0.1, 0.25, 0.5, 1, 2.5, 5],
});

// Total request count — enables rate() queries (requests per second)
const httpRequestsTotal = new client.Counter({
  name:       'http_requests_total',
  help:       'Total number of HTTP requests',
  labelNames: ['method', 'route', 'status_code'],
});

// Redis cache hit/miss counters — powers the Cache Hit Rate Grafana panel
const cacheHitsTotal   = new client.Counter({ name: 'cache_hits_total',   ... });
const cacheMissesTotal = new client.Counter({ name: 'cache_misses_total', ... });

module.exports = { client, httpRequestDuration, httpRequestsTotal,
                   cacheHitsTotal, cacheMissesTotal };""")

add_two_col_table(doc, [
    ("http_request_duration_seconds", "Histogram",
     "method, route, status_code",
     "Per-request latency. Buckets enable histogram_quantile() for p95/p99 latency panels."),
    ("http_requests_total",           "Counter",
     "method, route, status_code",
     "Total requests. rate() over 2m gives requests/second. Denominator for error rate."),
    ("cache_hits_total",              "Counter",
     "none",
     "Incremented on Redis cache hit. Combined with cache_misses_total for hit-rate panel."),
    ("cache_misses_total",            "Counter",
     "none",
     "Incremented on Redis cache miss (triggers MySQL query). Shows cache effectiveness."),
    ("process_cpu_seconds_total",     "Counter (default)",
     "none",
     "Node.js process CPU time — from collectDefaultMetrics()."),
    ("nodejs_heap_size_used_bytes",   "Gauge (default)",
     "none",
     "Node.js heap memory — early indicator of memory leaks."),
], header=["Metric Name", "Type", "Labels", "Purpose"])

add_subheading(doc, "27.2  HTTP Middleware in app/backend/src/app.js", level=2)
add_para(doc, (
    "An Express middleware intercepts every request/response pair and records duration and count "
    "with the correct route label. The middleware fires on res.on('finish') — after the response "
    "is sent — ensuring the measured duration includes all handler time (DB queries, Redis calls, "
    "JSON serialization)."
))
add_code_block(doc, """\
app.use((req, res, next) => {
  const end = httpRequestDuration.startTimer();   // starts the clock
  res.on('finish', () => {
    // Use Express route pattern (e.g. /api/todos/:id) not raw URL
    // This prevents high-cardinality labels from unique IDs like /api/todos/12345
    const rawRoute = req.route ? req.baseUrl + req.route.path : req.path;
    const route    = rawRoute.replace(/\\/$/, '') || '/';
    const labels   = { method: req.method, route, status_code: res.statusCode };
    end(labels);                    // records duration with labels
    httpRequestsTotal.inc(labels);  // increments request counter with labels
  });
  next();
});""")

add_note_box(doc,
    "Using req.route.path instead of req.path is essential for Prometheus. req.path would produce "
    "a unique label for every request with a numeric ID (e.g. /api/todos/1, /api/todos/2, ...), "
    "causing Prometheus cardinality explosion. req.route.path gives the template: /api/todos/:id.",
    "IMPORTANT"
)

add_subheading(doc, "27.3  /metrics Endpoint", level=2)
add_code_block(doc, """\
app.get('/metrics', async (_req, res) => {
  res.set('Content-Type', client.register.contentType);
  res.end(await client.register.metrics());
});""")
add_para(doc, (
    "The /metrics endpoint returns all registered metrics in Prometheus text exposition format "
    "(text/plain; version=0.0.4). Prometheus scrapes this endpoint every 15 seconds as configured "
    "in the ServiceMonitor. The endpoint is unauthenticated — in a stricter environment, it should "
    "be restricted to the monitoring namespace CIDR via NetworkPolicy or basic auth."
))
add_code_block(doc, """\
# Test the metrics endpoint locally:
kubectl port-forward -n todo deployment/todo-backend 3000:3000
curl http://localhost:3000/metrics

# Example output:
# HELP http_requests_total Total number of HTTP requests
# TYPE http_requests_total counter
# http_requests_total{method="GET",route="/api/todos",status_code="200"} 42
# http_requests_total{method="POST",route="/api/todos",status_code="200"} 7
#
# HELP http_request_duration_seconds Duration of HTTP requests in seconds
# TYPE http_request_duration_seconds histogram
# http_request_duration_seconds_bucket{le="0.025",...} 45
# ...
# http_request_duration_seconds_sum{...} 1.23
# http_request_duration_seconds_count{...} 49""")

add_subheading(doc, "27.4  Cache Hit/Miss Instrumentation in todo.controller.js", level=2)
add_code_block(doc, """\
const { cacheHitsTotal, cacheMissesTotal } = require('../metrics');
const CACHE_KEY = 'todos:all';
const CACHE_TTL = 60;   // seconds

exports.getTodos = async (req, res) => {
  const cached = await redis.get(CACHE_KEY);
  if (cached) {
    cacheHitsTotal.inc();                           // ← increments cache_hits_total
    return res.json(JSON.parse(cached));
  }
  cacheMissesTotal.inc();                           // ← increments cache_misses_total
  const [rows] = await req.db.execute("SELECT * FROM todos");
  await redis.set(CACHE_KEY, JSON.stringify(rows), 'EX', CACHE_TTL);
  res.json(rows);
};

// All write operations invalidate the cache:
exports.createTodo = async (req, res) => {
  // ... INSERT ...
  await req.redis.del(CACHE_KEY);   // cache-aside invalidation
  res.json({ message: "Todo created" });
};""")
add_para(doc, (
    "The cache-aside pattern means cache_misses_total increments on the first request after a "
    "write operation (or after the 60-second TTL expires). A sustained hit rate near 1.0 indicates "
    "the read-heavy workload is well-served by the cache. A hit rate near 0 indicates either "
    "very frequent writes, a Redis connection problem, or the first request after a cold start."
))

add_subheading(doc, "27.5  Useful PromQL Queries", level=2)
add_two_col_table(doc, [
    ("Request rate (req/s) by route",
     "sum(rate(http_requests_total[2m])) by (route)"),
    ("Server error rate (5xx %)",
     "100 * sum(rate(http_requests_total{status_code=~\"5..\"}[2m])) / sum(rate(http_requests_total[2m]))"),
    ("P95 response latency by route",
     "histogram_quantile(0.95, sum(rate(http_request_duration_seconds_bucket[2m])) by (le, route))"),
    ("Redis cache hit rate",
     "rate(cache_hits_total[2m]) / (rate(cache_hits_total[2m]) + rate(cache_misses_total[2m]))"),
    ("Backend pod count (ready)",
     "count(kube_pod_status_ready{namespace=\"todo\", pod=~\"todo-backend.*\", condition=\"true\"})"),
    ("Backend memory per pod (MB)",
     "sum(container_memory_working_set_bytes{namespace=\"todo\", pod=~\"todo-backend.*\"}) by (pod) / 1024 / 1024"),
    ("Backend CPU cores per pod",
     "sum(rate(container_cpu_usage_seconds_total{namespace=\"todo\", pod=~\"todo-backend.*\"}[2m])) by (pod)"),
    ("Todo CRUD operations by method",
     "sum(rate(http_requests_total{route=~\"/api/todos.*\"}[2m])) by (method)"),
], header=["What to measure", "PromQL"])

# ── Save ─────────────────────────────────────────────────────
out_path = "/home/vvdn/Desktop/todo-fullstack/docs/Fullstack_Deployment_using_Docker_Kubernetes_Terraform_Jenkins.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
